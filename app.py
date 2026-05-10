import hmac
import json
import os
import re
import tomllib

from collections.abc import Mapping

import streamlit as st
from streamlit.errors import StreamlitSecretNotFoundError
import base64
from io import BytesIO
import streamlit.components.v1 as components
import hashlib
import time
import unicodedata
from datetime import date, datetime, timedelta, timezone
from secrets import token_urlsafe
from pathlib import Path
import traceback
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from groq import Groq
from docx import Document
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import difflib
from passlib.context import CryptContext
from passlib.hash import bcrypt as _bcrypt_hash
from sqlalchemy import (
    JSON,
    Boolean,
    DateTime,
    ForeignKey,
    Integer,
    String,
    Text,
    create_engine,
    false as sql_false,
    inspect,
    select,
    text,
)
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import DeclarativeBase, Mapped, Session, mapped_column, relationship


DEFAULT_BRANDING: dict[str, object] = {
    "app_name": "Fluency",
    "tagline": "Capture expertise in a snap",
    "page_title": "Fluency",
    "page_icon": "🗣️",
    "logo_url": "",
    "logo_path": "",
    "primary_color": "#E1306C",
    "secondary_color": "#833AB4",
    "accent_color": "#FCAF45",
    "hide_powered_by": True,
    "access_gate_button": "Continue",
    "sign_in_button": "Sign in",
}

_BRANDING_KEYS = frozenset(DEFAULT_BRANDING.keys())


def _normalize_secret_value(value: object) -> str | None:
    if value is None:
        return None
    s = str(value).strip()
    return s or None


def _load_local_secrets() -> dict[str, object]:
    """Load local secrets from disk (dev convenience).

    Streamlit Cloud uses `st.secrets` and does not rely on these files.
    Locally, this allows password protection even if `.streamlit/secrets.toml`
    can't be created (e.g., name collision with an existing folder).
    """
    candidates = [
        os.path.join(".streamlit", "secrets.toml"),
        os.path.join(".streamlit", "secrets.local.toml"),
    ]
    for path in candidates:
        if not os.path.isfile(path):
            continue
        try:
            with open(path, "rb") as f:
                data = tomllib.load(f)
            return data if isinstance(data, dict) else {}
        except (OSError, tomllib.TOMLDecodeError):
            continue
    return {}


_LOCAL_SECRETS = _load_local_secrets()

# --- Logging + safe error UX (must be defined before auth) ---
_DEFAULT_LOG_PATH = os.path.join(os.getenv("TMPDIR") or os.getenv("TEMP") or "/tmp", "fluency-app.log")
APP_LOG_PATH = os.getenv("APP_LOG_PATH") or _DEFAULT_LOG_PATH

# Persistent login (full page refresh / new Streamlit session). Requires SESSION_COOKIE_SECRET in secrets.
SESSION_COOKIE_NAME = "fluency_auth"


def log_exception(ex: Exception, *, context: str) -> None:
    """Write exception details to a local log file for support/debugging."""
    try:
        os.makedirs(os.path.dirname(APP_LOG_PATH), exist_ok=True)
        with open(APP_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(f"\n[{datetime.now(timezone.utc).isoformat()}] {context}\n")
            f.write("".join(traceback.format_exception(type(ex), ex, ex.__traceback__)))
    except Exception:
        # Never let logging break the app UX.
        try:
            st.session_state["_last_exception_context"] = context
            st.session_state["_last_exception_text"] = "".join(
                traceback.format_exception(type(ex), ex, ex.__traceback__)
            )
        except Exception:
            pass


def show_busy_error(ex: Exception | None = None, *, context: str = "Unhandled error") -> None:
    if ex is not None:
        log_exception(ex, context=context)
    st.error("Something went wrong. Please try again. If this continues, contact support.")

    # Optional debug details (admin-only or DEBUG_ERRORS=true). Safe: never shows secrets unless you paste them.
    try:
        debug = str(_secret_or_env("DEBUG_ERRORS") or "").strip().lower() in {"1", "true", "yes", "on"}
        role = str((st.session_state.get("auth_user") or {}).get("role") or "")
        if debug or role == "admin":
            with st.expander("Error details (admin/debug)", expanded=False):
                st.caption(f"Context: {context}")
                if ex is not None:
                    st.exception(ex)
                else:
                    st.write(st.session_state.get("_last_exception_text") or "No exception captured.")
                st.caption(f"Log file path: {APP_LOG_PATH}")
    except Exception:
        pass


def _secret_or_env(name: str) -> str | None:
    """Read a scalar secret from secrets.toml when present; fall back to the same-named env var."""
    try:
        if name in st.secrets:
            return _normalize_secret_value(st.secrets[name])
    except (StreamlitSecretNotFoundError, OSError, PermissionError):
        pass
    if name in _LOCAL_SECRETS:
        return _normalize_secret_value(_LOCAL_SECRETS.get(name))
    return _normalize_secret_value(os.getenv(name))


def _secret_first(*keys: str) -> str | None:
    """Return the first non-empty secret/env match (Streamlit Cloud keys vary by casing)."""
    for k in keys:
        v = _secret_or_env(k)
        if v:
            return v
    return None


def _database_url() -> str:
    """Database connection string.

    - SaaS/prod: set DATABASE_URL (e.g. Postgres).
    - Local dev: defaults to a local sqlite file.
    """
    url = _secret_first(
        "DATABASE_URL",
        "database_url",
        "database_URL",
    )
    if url:
        return url.strip()
    return "sqlite+pysqlite:///./fluency.db"


def _database_backend_label() -> str:
    """Human-readable DB backend for smoke tests (no credentials)."""
    u = _database_url().strip().lower()
    if u.startswith("postgresql") or u.startswith("postgres"):
        return "PostgreSQL"
    if u.startswith("sqlite"):
        return "SQLite (dev only)"
    return "Database"


def _coerce_int(value: object, default: int) -> int:
    try:
        if value is None:
            return default
        return int(str(value).strip())
    except Exception:
        return default


def _normalize_email(email: str) -> str:
    s = (email or "").strip().lower()
    return unicodedata.normalize("NFC", s)


def _normalize_password_input(password: str) -> str:
    """Strip accidental spaces/newlines from pasted passwords (signup + login)."""
    return (password or "").strip()


_PWD = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")


def _hash_password(password: str) -> str:
    return _PWD.hash(_normalize_password_input(password))


def _verify_password(password: str, password_hash: str) -> bool:
    """Verify password; try trimmed input first, then raw (legacy hashes from before strip fix)."""
    ph = (password_hash or "").strip()
    if not ph:
        return False
    try:
        if _PWD.verify(_normalize_password_input(password), ph):
            return True
        raw = password or ""
        if raw != _normalize_password_input(password):
            return bool(_PWD.verify(raw, ph))
        return False
    except Exception:
        return False


def _verify_and_migrate_password(
    *,
    tenant_id: int | None,
    user_id: int | None,
    password: str,
    password_hash: str,
) -> bool:
    """Verify password, and migrate legacy bcrypt hashes to pbkdf2_sha256 on success."""
    pw = _normalize_password_input(password)
    ph = password_hash or ""

    # First try current scheme.
    if _verify_password(pw, ph):
        return True

    # Legacy support: bcrypt hashes start with $2a$ / $2b$ / $2y$.
    if ph.startswith("$2a$") or ph.startswith("$2b$") or ph.startswith("$2y$"):
        try:
            # bcrypt only uses first 72 bytes; avoid passlib throwing.
            pw72 = _normalize_password_input(pw).encode("utf-8")[:72].decode("utf-8", "ignore")
            ok = _bcrypt_hash.verify(pw72, ph)
        except Exception:
            ok = False
        if not ok:
            return False

        # Migrate to pbkdf2_sha256 (best-effort; never block login if migration fails).
        if tenant_id is not None and user_id is not None:
            try:
                _init_db()
                with _db() as s:
                    u = s.execute(
                        select(User).where(User.tenant_id == int(tenant_id), User.id == int(user_id))
                    ).scalar_one_or_none()
                    if u is not None:
                        u.password_hash = _hash_password(pw)
                        s.commit()
            except Exception:
                pass
        return True

    return False


class Base(DeclarativeBase):
    pass


class Tenant(Base):
    __tablename__ = "tenants"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    slug: Mapped[str] = mapped_column(String(120), unique=True, index=True)
    name: Mapped[str] = mapped_column(String(200))
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))

    users: Mapped[list["User"]] = relationship(back_populates="tenant")


class User(Base):
    __tablename__ = "users"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    tenant_id: Mapped[int] = mapped_column(ForeignKey("tenants.id"), index=True)
    email: Mapped[str] = mapped_column(String(320), index=True)
    # Text: room for any passlib scheme; avoids rare truncation on some DBs
    password_hash: Mapped[str] = mapped_column(Text)
    role: Mapped[str] = mapped_column(String(50), default="member")  # admin | reviewer | member
    is_active: Mapped[bool] = mapped_column(Boolean, default=True)
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))
    # Self-signup: False until magic-link verified (no SMTP — user opens link from dashboard message).
    email_verified: Mapped[bool] = mapped_column(Boolean, default=True)
    verification_token: Mapped[str | None] = mapped_column(String(128), nullable=True)
    verification_expires_at: Mapped[datetime | None] = mapped_column(DateTime(timezone=True), nullable=True)

    tenant: Mapped["Tenant"] = relationship(back_populates="users")


class TenantSettings(Base):
    __tablename__ = "tenant_settings"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    tenant_id: Mapped[int] = mapped_column(ForeignKey("tenants.id"), unique=True, index=True)
    audience: Mapped[str] = mapped_column(String(500), default="")
    tools_used: Mapped[str] = mapped_column(String(700), default="")
    compliance_standard: Mapped[str] = mapped_column(String(120), default="")
    tone: Mapped[str] = mapped_column(String(120), default="Professional")
    updated_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class SopDoc(Base):
    __tablename__ = "sop_docs"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    tenant_id: Mapped[int] = mapped_column(ForeignKey("tenants.id"), index=True)
    title: Mapped[str] = mapped_column(String(300), default="SOP")
    template_name: Mapped[str] = mapped_column(String(120), default="IT SOP")
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))
    created_by_user_id: Mapped[int | None] = mapped_column(ForeignKey("users.id"), nullable=True)

    versions: Mapped[list["SopVersion"]] = relationship(back_populates="doc", cascade="all, delete-orphan")


class SopVersion(Base):
    __tablename__ = "sop_versions"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    tenant_id: Mapped[int] = mapped_column(ForeignKey("tenants.id"), index=True)
    sop_doc_id: Mapped[int] = mapped_column(ForeignKey("sop_docs.id"), index=True)
    version: Mapped[int] = mapped_column(Integer, default=1)
    status: Mapped[str] = mapped_column(String(30), default="draft")  # draft | in_review | approved
    label: Mapped[str] = mapped_column(String(300), default="")
    source: Mapped[str] = mapped_column(String(30), default="generated")  # generated | revised | edited
    sop_text: Mapped[str] = mapped_column(Text)
    sop_sha256: Mapped[str] = mapped_column(String(64), index=True)
    change_note: Mapped[str] = mapped_column(String(500), default="")
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))
    created_by_user_id: Mapped[int | None] = mapped_column(ForeignKey("users.id"), nullable=True)
    approved_at: Mapped[datetime | None] = mapped_column(DateTime(timezone=True), nullable=True)
    approved_by_user_id: Mapped[int | None] = mapped_column(ForeignKey("users.id"), nullable=True)

    doc: Mapped["SopDoc"] = relationship(back_populates="versions")


class Feedback(Base):
    __tablename__ = "feedback"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    tenant_id: Mapped[int] = mapped_column(ForeignKey("tenants.id"), index=True)
    ts: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))
    rating: Mapped[str] = mapped_column(String(10))  # up | down
    reason: Mapped[str] = mapped_column(Text, default="")
    template_name: Mapped[str] = mapped_column(String(120), default="")
    strictness: Mapped[str] = mapped_column(String(30), default="")
    tone: Mapped[str] = mapped_column(String(60), default="")
    compliance_standard: Mapped[str] = mapped_column(String(120), default="")
    audience: Mapped[str] = mapped_column(String(500), default="")
    tools_used: Mapped[str] = mapped_column(String(700), default="")
    include_definitions: Mapped[bool] = mapped_column(Boolean, default=False)
    include_safety_compliance: Mapped[bool] = mapped_column(Boolean, default=False)
    include_records: Mapped[bool] = mapped_column(Boolean, default=False)
    include_checklist: Mapped[bool] = mapped_column(Boolean, default=False)
    model: Mapped[str] = mapped_column(String(120), default="")
    temperature: Mapped[int] = mapped_column(Integer, default=0)  # store temp*100 for simplicity
    notes_chars: Mapped[int] = mapped_column(Integer, default=0)
    sop_sha256: Mapped[str] = mapped_column(String(64), default="")
    source: Mapped[str] = mapped_column(String(30), default="")
    meta: Mapped[dict] = mapped_column(JSON, default=dict)


class ManualDoc(Base):
    __tablename__ = "manual_docs"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    tenant_id: Mapped[int] = mapped_column(ForeignKey("tenants.id"), index=True)
    name: Mapped[str] = mapped_column(String(260))
    sha256: Mapped[str] = mapped_column(String(64), index=True)
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))

    chunks: Mapped[list["ManualChunk"]] = relationship(back_populates="doc", cascade="all, delete-orphan")


class ManualChunk(Base):
    __tablename__ = "manual_chunks"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    tenant_id: Mapped[int] = mapped_column(ForeignKey("tenants.id"), index=True)
    manual_doc_id: Mapped[int] = mapped_column(ForeignKey("manual_docs.id"), index=True)
    chunk_index: Mapped[int] = mapped_column(Integer)
    text: Mapped[str] = mapped_column(Text)

    doc: Mapped["ManualDoc"] = relationship(back_populates="chunks")


class TenantQuota(Base):
    __tablename__ = "tenant_quotas"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    tenant_id: Mapped[int] = mapped_column(ForeignKey("tenants.id"), unique=True, index=True)
    generations_per_day: Mapped[int] = mapped_column(Integer, default=50)
    transcriptions_per_day: Mapped[int] = mapped_column(Integer, default=50)
    vision_analyses_per_day: Mapped[int] = mapped_column(Integer, default=50)
    updated_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


class DailyUsage(Base):
    __tablename__ = "daily_usage"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    tenant_id: Mapped[int] = mapped_column(ForeignKey("tenants.id"), index=True)
    day: Mapped[str] = mapped_column(String(10), index=True)  # YYYY-MM-DD
    action: Mapped[str] = mapped_column(String(40), index=True)  # generate | transcribe | vision
    count: Mapped[int] = mapped_column(Integer, default=0)
    updated_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=lambda: datetime.now(timezone.utc))


@st.cache_resource
def _engine(database_url: str):
    """One engine per DATABASE_URL. Uncached URL would pin SQLite forever when secrets load late."""
    url = (database_url or "").strip()
    if not url:
        url = "sqlite+pysqlite:///./fluency.db"
    # Avoid noisy check_same_thread issues for sqlite under Streamlit.
    connect_args = {"check_same_thread": False} if url.startswith("sqlite") else {}
    # Neon/serverless Postgres often closes idle connections — recycle pools.
    kw = dict(pool_pre_ping=True, connect_args=connect_args)
    if url.startswith("postgresql") or url.startswith("postgres"):
        kw["pool_recycle"] = 300
    return create_engine(url, **kw)


def _migrate_user_verification_columns(engine) -> None:
    """Add email verification columns to existing deployments (Neon/SQLite)."""
    try:
        insp = inspect(engine)
        cols = {c["name"] for c in insp.get_columns("users")}
    except Exception:
        return
    dialect = engine.dialect.name
    with engine.begin() as conn:
        if "email_verified" not in cols:
            if dialect == "sqlite":
                conn.execute(text("ALTER TABLE users ADD COLUMN email_verified BOOLEAN DEFAULT 1"))
            else:
                conn.execute(text("ALTER TABLE users ADD COLUMN email_verified BOOLEAN DEFAULT TRUE"))
        if "verification_token" not in cols:
            conn.execute(text("ALTER TABLE users ADD COLUMN verification_token VARCHAR(128)"))
        if "verification_expires_at" not in cols:
            if dialect == "sqlite":
                conn.execute(text("ALTER TABLE users ADD COLUMN verification_expires_at TIMESTAMP"))
            else:
                conn.execute(text("ALTER TABLE users ADD COLUMN verification_expires_at TIMESTAMP WITH TIME ZONE"))


def _init_db() -> None:
    eng = _engine(_database_url())
    Base.metadata.create_all(eng)
    _migrate_user_verification_columns(eng)


def _db() -> Session:
    return Session(_engine(_database_url()))


def _ensure_tenant_settings_and_quota(s: Session, *, tenant_id: int) -> None:
    """Ensure TenantSettings and TenantQuota rows exist for a tenant."""
    ts = s.execute(select(TenantSettings).where(TenantSettings.tenant_id == tenant_id)).scalar_one_or_none()
    if ts is None:
        s.add(TenantSettings(tenant_id=tenant_id))
    tq = s.execute(select(TenantQuota).where(TenantQuota.tenant_id == tenant_id)).scalar_one_or_none()
    if tq is None:
        s.add(
            TenantQuota(
                tenant_id=tenant_id,
                generations_per_day=_coerce_int(_secret_first("DEFAULT_GENERATIONS_PER_DAY"), 50),
                transcriptions_per_day=_coerce_int(_secret_first("DEFAULT_TRANSCRIPTIONS_PER_DAY"), 50),
                vision_analyses_per_day=_coerce_int(_secret_first("DEFAULT_VISION_PER_DAY"), 50),
            )
        )


def _ensure_bootstrap_admin() -> None:
    """Create bootstrap admin from Streamlit secrets / env if that user does not exist yet.

    Reads (first non-empty wins for alternate spellings):
    - BOOTSTRAP_ADMIN_EMAIL / bootstrap_admin_email
    - BOOTSTRAP_ADMIN_PASSWORD / bootstrap_admin_password
    - BOOTSTRAP_TENANT_SLUG / bootstrap_tenant_slug (default: default)
    - BOOTSTRAP_TENANT_NAME / bootstrap_tenant_name (default: Default Tenant)

    Behavior:
    - If email or password is missing, skip (interactive first-admin form still works).
    - If a user with BOOTSTRAP_ADMIN_EMAIL already exists, skip (idempotent).
    - Otherwise find tenant by slug or create it, then create the admin user on Neon/Postgres.

    This runs after `_init_db()` so tables exist on a fresh Neon database.
    """
    email = _normalize_email(
        _secret_first("BOOTSTRAP_ADMIN_EMAIL", "bootstrap_admin_email") or ""
    )
    pwd = _secret_first("BOOTSTRAP_ADMIN_PASSWORD", "bootstrap_admin_password") or ""

    if not email or not pwd:
        return

    slug = (_secret_first("BOOTSTRAP_TENANT_SLUG", "bootstrap_tenant_slug") or "default").strip()
    name = (_secret_first("BOOTSTRAP_TENANT_NAME", "bootstrap_tenant_name") or "Default Tenant").strip()

    try:
        with _db() as s:
            existing = s.execute(select(User).where(User.email == email)).scalar_one_or_none()
            if existing is not None:
                return

            t = s.execute(select(Tenant).where(Tenant.slug == slug)).scalar_one_or_none()
            if t is None:
                t = Tenant(slug=slug, name=name or slug)
                s.add(t)
                s.flush()

            _ensure_tenant_settings_and_quota(s, tenant_id=t.id)

            u = User(
                tenant_id=t.id,
                email=email,
                password_hash=_hash_password(pwd),
                role="admin",
                is_active=True,
                email_verified=True,
            )
            s.add(u)
            s.commit()
    except Exception as ex:
        log_exception(ex, context="Bootstrap admin from secrets")


def _login_required() -> bool:
    return not _coerce_bool(_secret_or_env("AUTH_DISABLED"), False)


def _self_signup_allowed() -> bool:
    """SaaS self-service ‘Create workspace’. Set SELF_SIGNUP_ENABLED=false in secrets to disable.

    When the secret is **unset**, signup is **on** (default for SaaS). Using `_coerce_bool(None, True)`
    is wrong: `str(None) == "none"` would incorrectly turn signup off.
    """
    raw = _secret_first("SELF_SIGNUP_ENABLED", "ALLOW_SELF_SIGNUP")
    if raw is None or str(raw).strip() == "":
        return True
    return _coerce_bool(raw, False)


def _sanitize_tenant_slug(raw: str) -> str:
    s = (raw or "").strip().lower()
    s = re.sub(r"[^a-z0-9]+", "-", s)
    s = re.sub(r"-+", "-", s).strip("-")
    return (s[:118] if s else "workspace")


def _signup_requires_email_verification() -> bool:
    """Self-signup must verify via magic link (no SMTP). Set REQUIRE_EMAIL_VERIFICATION=false to skip."""
    raw = _secret_first("REQUIRE_EMAIL_VERIFICATION", "EMAIL_VERIFICATION_REQUIRED")
    if raw is None or str(raw).strip() == "":
        return True
    return _coerce_bool(raw, True)


def _app_base_url() -> str:
    return (_secret_first("APP_BASE_URL", "PUBLIC_URL") or "").rstrip("/")


def _verification_url_for_token(token: str) -> str:
    base = (_app_base_url() or "").strip().rstrip("/")
    tok = (token or "").strip()
    if not tok or not base:
        return ""
    return f"{base}?verify_token={tok}"


def _consume_verification_token(raw_token: str) -> bool:
    """Single-use magic link token (stored server-side)."""
    tok = (raw_token or "").strip()
    if not tok:
        return False
    if "verify_token=" in tok:
        frag = tok.split("verify_token=", 1)[1].split("&")[0].split("#")[0].strip()
        if frag:
            tok = frag
    _init_db()
    now = datetime.now(timezone.utc)
    try:
        with _db() as s:
            u = s.execute(select(User).where(User.verification_token == tok)).scalars().first()
            if u is None:
                return False
            exp = u.verification_expires_at
            if exp is not None and exp.tzinfo is None:
                exp = exp.replace(tzinfo=timezone.utc)
            if exp is not None and exp < now:
                return False
            u.email_verified = True
            u.verification_token = None
            u.verification_expires_at = None
            s.commit()
        return True
    except Exception as ex:
        log_exception(ex, context="Consume verification token")
        return False


def _issue_new_verification_token(*, email: str) -> tuple[bool, str]:
    """Regenerate token for an unverified user (still no SMTP — user must open link)."""
    email_n = _normalize_email(email)
    if not email_n:
        return False, "Enter your email address."
    hrs = _coerce_int(_secret_first("VERIFICATION_LINK_EXPIRY_HOURS"), 48)
    if hrs < 1:
        hrs = 48
    _init_db()
    try:
        with _db() as s:
            u = s.execute(select(User).where(User.email == email_n)).scalars().first()
            if u is None:
                return False, "No account found for that email."
            if getattr(u, "email_verified", True) is True:
                return False, "That email is already verified. Sign in."
            tok = token_urlsafe(48)
            u.verification_token = tok
            u.verification_expires_at = datetime.now(timezone.utc) + timedelta(hours=hrs)
            s.commit()
        url = _verification_url_for_token(tok)
        if url:
            return True, url
        return True, f"(Set APP_BASE_URL for a clickable link.) Append to your app URL: ?verify_token={tok}"
    except Exception as ex:
        log_exception(ex, context="Issue verification token")
        return False, "Could not issue a new link. Try again."


def register_workspace(
    *,
    slug_raw: str,
    display_name: str,
    email: str,
    password: str,
) -> tuple[bool, str, dict | None]:
    """Create a new tenant + first admin (self-service signup).

    Returns (ok, message, extra). When email verification is required, extra contains
    verification_url / token for the magic-link flow (no SMTP).
    """
    slug = _sanitize_tenant_slug(slug_raw)
    if len(slug) < 2:
        return False, "Workspace URL must be at least 2 characters (letters or numbers).", None
    name = (display_name or "").strip() or slug.replace("-", " ").title()
    email_n = _normalize_email(email)
    if not email_n:
        return False, "Enter a valid email address.", None
    pw_n = _normalize_password_input(password)
    if len(pw_n) < 8:
        return False, "Password must be at least 8 characters.", None

    verify_on = _signup_requires_email_verification()
    hrs = _coerce_int(_secret_first("VERIFICATION_LINK_EXPIRY_HOURS"), 48)
    if hrs < 1:
        hrs = 48
    vtok = token_urlsafe(48) if verify_on else None
    vexp = (datetime.now(timezone.utc) + timedelta(hours=hrs)) if verify_on else None

    _init_db()
    try:
        with _db() as s:
            if s.execute(select(User.id).where(User.email == email_n)).scalar_one_or_none() is not None:
                return False, "That email is already registered. Sign in instead.", None
            if s.execute(select(Tenant.id).where(Tenant.slug == slug)).scalar_one_or_none() is not None:
                return False, "That workspace URL is already taken. Choose another.", None

            t = Tenant(slug=slug, name=name)
            s.add(t)
            s.flush()
            _ensure_tenant_settings_and_quota(s, tenant_id=t.id)
            s.add(
                User(
                    tenant_id=t.id,
                    email=email_n,
                    password_hash=_hash_password(pw_n),
                    role="admin",
                    is_active=True,
                    email_verified=(False if verify_on else True),
                    verification_token=vtok,
                    verification_expires_at=vexp,
                )
            )
            s.commit()
    except IntegrityError:
        return False, "Workspace URL or email conflict. Try a different URL or sign in.", None
    except Exception as ex:
        log_exception(ex, context="Register workspace")
        return False, "Could not create workspace. Try again or contact support.", None

    if verify_on and vtok:
        url = _verification_url_for_token(vtok)
        msg = (
            "Workspace created. Verify your email before signing in — open the magic link below "
            "(no SMTP in-app; you can paste the link into your browser or send it to yourself from your mailbox)."
        )
        extra = {
            "verification_url": url,
            "verification_token": vtok,
            "expires_hours": hrs,
        }
        return True, msg, extra

    return True, "Workspace created. Sign in with your email and password.", None


def _current_user() -> dict | None:
    return st.session_state.get("auth_user")


def _session_cookie_secret_explicit() -> str | None:
    return _normalize_secret_value(
        _secret_first("SESSION_COOKIE_SECRET", "FLUENCY_SESSION_SECRET")
    )


def _session_cookie_secret_effective() -> str | None:
    """Secret for signing session cookies: explicit `SESSION_COOKIE_SECRET`, else auto local file.

    Streamlit Cloud must set `SESSION_COOKIE_SECRET` (filesystem is not durable).
    Local dev: writes `.streamlit/.session_cookie_secret` once so refresh works without manual setup.
    """
    ex = _session_cookie_secret_explicit()
    if ex:
        return ex
    path = os.path.join(".streamlit", ".session_cookie_secret")
    try:
        if os.path.isfile(path):
            with open(path, encoding="utf-8") as f:
                s = _normalize_secret_value(f.read())
                if s:
                    return s
        os.makedirs(os.path.dirname(path), exist_ok=True)
        tok = token_urlsafe(48)
        with open(path, "w", encoding="utf-8") as f:
            f.write(tok)
            f.flush()
            try:
                os.fsync(f.fileno())
            except Exception:
                pass
        return tok
    except OSError:
        return None


def _session_cookie_max_age_seconds() -> int:
    days = _coerce_int(_secret_first("SESSION_COOKIE_DAYS", "SESSION_MAX_AGE_DAYS"), 30)
    return max(1, min(days, 366)) * 86400


def _sign_session_token(user_id: int) -> str:
    secret = _session_cookie_secret_effective()
    if not secret:
        return ""
    exp = int(time.time()) + _session_cookie_max_age_seconds()
    payload = json.dumps({"u": int(user_id), "exp": exp}, separators=(",", ":")).encode("utf-8")
    b64 = base64.urlsafe_b64encode(payload).decode("ascii").rstrip("=")
    msg = f"v1.{b64}"
    sig = hmac.new(secret.encode("utf-8"), msg.encode("utf-8"), hashlib.sha256).hexdigest()
    return f"{msg}.{sig}"


def _verify_session_token(token: str) -> int | None:
    secret = _session_cookie_secret_effective()
    if not secret or not (token or "").strip():
        return None
    parts = token.strip().split(".")
    if len(parts) != 3 or parts[0] != "v1":
        return None
    msg = f"{parts[0]}.{parts[1]}"
    sig = parts[2]
    try:
        expected = hmac.new(secret.encode("utf-8"), msg.encode("utf-8"), hashlib.sha256).hexdigest()
        if not hmac.compare_digest(expected, sig):
            return None
        pad = "=" * (-len(parts[1]) % 4)
        data = json.loads(base64.urlsafe_b64decode(parts[1] + pad))
        if int(data.get("exp", 0)) < int(time.time()):
            return None
        uid = data.get("u")
        return int(uid) if uid is not None else None
    except Exception:
        return None


def _cookie_header_session_value() -> str | None:
    """Parse Cookie header (often populated when st.context.cookies is empty)."""
    try:
        h = getattr(st.context, "headers", None)
        if h is None:
            return None
        raw = h.get("cookie") or h.get("Cookie")
        if not raw:
            return None
        target = SESSION_COOKIE_NAME.lower()
        for segment in str(raw).split(";"):
            segment = segment.strip()
            if "=" not in segment:
                continue
            k, v = segment.split("=", 1)
            if k.strip().lower() != target:
                continue
            out = v.strip().strip('"')
            try:
                from urllib.parse import unquote

                return unquote(out)
            except Exception:
                return out
    except Exception:
        pass
    return None


def _session_cookie_from_st_context_dict() -> str | None:
    try:
        ctx = getattr(st, "context", None)
        if ctx is None:
            return None
        cookies = getattr(ctx, "cookies", None)
        if not cookies:
            return None
        for name in (SESSION_COOKIE_NAME, SESSION_COOKIE_NAME.lower()):
            if hasattr(cookies, "get"):
                v = cookies.get(name)
                if v:
                    return str(v)
        if hasattr(cookies, "items"):
            for k, v in cookies.items():
                if str(k).lower() == SESSION_COOKIE_NAME.lower() and v:
                    return str(v)
        return None
    except Exception:
        return None


def _session_cookie_token_raw() -> str | None:
    """Resolve session token: HTTP Cookie header, then st.context.cookies, then CookieManager."""
    t = _cookie_header_session_value()
    if t:
        return t
    t = _session_cookie_from_st_context_dict()
    if t:
        return t
    mgr = _cookie_manager_singleton()
    if mgr:
        try:
            mgr.get_all(key="fluency_ck_sync")
            v = mgr.get(SESSION_COOKIE_NAME)
            if v:
                return str(v)
        except Exception:
            pass
    return None


def _request_is_https() -> bool:
    try:
        h = getattr(st.context, "headers", None)
        if h is None:
            return False
        proto = h.get("x-forwarded-proto") or h.get("X-Forwarded-Proto")
        return str(proto or "").lower() == "https"
    except Exception:
        return False


def _cookie_manager_singleton():
    k = "_fluency_cookie_mgr"
    if k not in st.session_state:
        try:
            import extra_streamlit_components as stx

            st.session_state[k] = stx.CookieManager(key="fluency_auth_ck")
        except Exception:
            st.session_state[k] = None
    return st.session_state[k]


def _persist_session_cookie_js_fallback(token: str, max_age: int) -> None:
    safe = json.dumps(token)
    components.html(
        f"""<script>
(function() {{
  const t = {safe};
  let s = "{SESSION_COOKIE_NAME}=" + encodeURIComponent(t) + "; path=/; max-age=" + {int(max_age)} + "; SameSite=Lax";
  if (location.protocol === "https:") s += "; Secure";
  function apply(doc) {{
    try {{ doc.cookie = s; }} catch (e) {{}}
  }}
  apply(document);
  try {{
    if (window.parent && window.parent !== window && window.parent.document) {{
      apply(window.parent.document);
    }}
  }} catch (e) {{}}
}})();
</script>""",
        height=0,
        width=0,
    )


def _clear_session_cookie_js_fallback() -> None:
    components.html(
        f"""<script>
try {{
  document.cookie = "{SESSION_COOKIE_NAME}=; path=/; max-age=0; SameSite=Lax";
}} catch (e) {{}}
</script>""",
        height=0,
        width=0,
    )


def _persist_session_cookie(user_id: int) -> None:
    if not _session_cookie_secret_effective():
        return
    tok = _sign_session_token(user_id)
    if not tok:
        return
    ma = _session_cookie_max_age_seconds()
    mgr = _cookie_manager_singleton()
    sec = True if _request_is_https() else None
    if mgr:
        try:
            mgr.set(
                SESSION_COOKIE_NAME,
                tok,
                path="/",
                max_age=float(ma),
                same_site="lax",
                secure=sec,
                key="persist_ck",
            )
        except Exception as ex:
            log_exception(ex, context="CookieManager.set session")
        # Always mirror via JS: iframe + top-frame so the browser sends Cookie on the next request.
        _persist_session_cookie_js_fallback(tok, ma)
    else:
        _persist_session_cookie_js_fallback(tok, ma)


def _clear_persistent_auth_cookie() -> None:
    mgr = _cookie_manager_singleton()
    if mgr:
        try:
            mgr.delete(SESSION_COOKIE_NAME, key="del_ck")
        except Exception:
            pass
    _clear_session_cookie_js_fallback()


def _restore_auth_from_signed_cookie() -> None:
    if st.session_state.get("auth_user"):
        return
    if not _session_cookie_secret_effective():
        return
    raw = _session_cookie_token_raw()
    if not raw:
        return
    uid = _verify_session_token(raw)
    if uid is None:
        _clear_persistent_auth_cookie()
        return
    try:
        with _db() as s:
            u = s.get(User, uid)
            if u is None or u.is_active is False:
                _clear_persistent_auth_cookie()
                return
            if getattr(u, "email_verified", True) is False:
                _clear_persistent_auth_cookie()
                return
            st.session_state.auth_user = {
                "user_id": u.id,
                "tenant_id": u.tenant_id,
                "email": u.email,
                "role": u.role,
            }
    except Exception as ex:
        log_exception(ex, context="Restore session from cookie")


def _require_auth_ui(brand: dict[str, object]) -> None:
    """Login screen + optional bootstrap setup + self-service signup."""
    if not _login_required():
        st.session_state.auth_user = {
            "user_id": None,
            "tenant_id": None,
            "email": "anonymous",
            "role": "admin",
        }
        return

    _init_db()
    _ensure_bootstrap_admin()

    user = _current_user()
    if user:
        return

    # Magic-link verification (?verify_token=...) — no SMTP; user opens link from on-screen instructions.
    try:
        qp = st.query_params
        if "verify_token" in qp:
            raw_tok = qp.get("verify_token")
            if isinstance(raw_tok, list):
                raw_tok = raw_tok[0]
            msg: str | None = None
            if raw_tok:
                if _consume_verification_token(str(raw_tok)):
                    msg = "ok"
                else:
                    msg = "err"
            try:
                del st.query_params["verify_token"]
            except Exception:
                pass
            # Rerun only when we actually handled a token — unconditional rerun caused blank/spin pages
            # when the query param could not be cleared or was empty.
            if msg:
                st.session_state["_verify_flash"] = msg
                st.rerun()
    except Exception as ex:
        log_exception(ex, context="Verify token from URL")

    vf = st.session_state.pop("_verify_flash", None)
    if vf == "ok":
        st.success("Email verified. You can sign in below.")
    elif vf == "err":
        st.error("Invalid or expired verification link.")

    st.title(str(brand.get("app_name") or DEFAULT_BRANDING["app_name"]))
    st.caption("Sign in or create your organization workspace.")

    # Optional: set SHOW_DB_HINT=true in secrets during redeploy smoke tests (PostgreSQL vs SQLite).
    try:
        hint_raw = _secret_first("SHOW_DB_HINT", "SHOW_DEPLOYMENT_INFO")
        if hint_raw is not None and str(hint_raw).strip() != "" and _coerce_bool(hint_raw, False):
            st.caption(f"Deployment DB: **{_database_backend_label()}**")
    except Exception:
        pass

    with _db() as s:
        any_user = s.execute(select(User.id).limit(1)).scalar_one_or_none()

    signup_ok = _self_signup_allowed()

    # No users yet and signup disabled: keep manual one-time bootstrap (no secrets).
    if any_user is None and not signup_ok:
        st.warning("No admin exists yet. Create the first admin (or set bootstrap secrets / enable self-signup).")
        with st.form("bootstrap_admin"):
            tenant_slug = st.text_input("Tenant slug", value="default")
            tenant_name = st.text_input("Tenant name", value="Default Tenant")
            email = st.text_input("Admin email")
            pwd = st.text_input("Admin password", type="password")
            if st.form_submit_button("Create admin"):
                email_n = _normalize_email(email)
                if not tenant_slug.strip() or not email_n or not pwd:
                    st.error("Tenant slug, email, and password are required.")
                else:
                    try:
                        with _db() as s:
                            t = Tenant(slug=tenant_slug.strip(), name=tenant_name.strip() or tenant_slug.strip())
                            s.add(t)
                            s.flush()
                            _ensure_tenant_settings_and_quota(s, tenant_id=t.id)
                            u = User(
                                tenant_id=t.id,
                                email=email_n,
                                password_hash=_hash_password(pwd),
                                role="admin",
                                is_active=True,
                                email_verified=True,
                            )
                            s.add(u)
                            s.commit()
                        st.success("Admin created. Please sign in.")
                    except Exception as e:
                        show_busy_error(e, context="Bootstrap admin")
        st.stop()

    tab_signin, tab_signup = st.tabs(["Sign in", "Create workspace"])

    with tab_signin:
        sign_btn = str(brand.get("sign_in_button") or "Sign in").strip() or "Sign in"
        with st.form("login"):
            email = st.text_input("Email", key="login_email")
            pwd = st.text_input("Password", type="password", key="login_pwd")
            if st.form_submit_button(sign_btn):
                try:
                    email_n = _normalize_email(email)
                    pwd_try = _normalize_password_input(pwd)
                    with _db() as s:
                        u = s.execute(select(User).where(User.email == email_n)).scalars().first()
                        debug_login = _secret_first("DEBUG_LOGIN")
                        show_debug = debug_login is not None and str(debug_login).strip() != "" and _coerce_bool(debug_login, False)

                        if u is None:
                            st.error("Invalid email or password.")
                            if show_debug:
                                with st.expander("Login debug (remove DEBUG_LOGIN from secrets later)"):
                                    st.write({"reason": "no_user_for_email", "email": email_n})
                        elif u.is_active is False:
                            st.error("This account is disabled. Contact your administrator.")
                        elif getattr(u, "email_verified", True) is False:
                            st.error(
                                "Email not verified yet. Open your verification link first, "
                                "or use **Resend link** below."
                            )
                        elif not u.password_hash or not str(u.password_hash).strip():
                            st.error("Account error: missing password hash. Reset password via admin or bootstrap.")
                            log_exception(
                                RuntimeError("User has empty password_hash"),
                                context=f"Login empty hash user_id={u.id}",
                            )
                        elif not _verify_and_migrate_password(
                            tenant_id=u.tenant_id,
                            user_id=u.id,
                            password=pwd_try,
                            password_hash=u.password_hash,
                        ):
                            st.error("Invalid email or password.")
                            if show_debug:
                                with st.expander("Login debug (remove DEBUG_LOGIN from secrets later)"):
                                    st.write(
                                        {
                                            "reason": "password_mismatch",
                                            "email": email_n,
                                            "user_id": u.id,
                                            "hash_prefix": (u.password_hash or "")[:20],
                                        }
                                    )
                        else:
                            st.session_state.auth_user = {
                                "user_id": u.id,
                                "tenant_id": u.tenant_id,
                                "email": u.email,
                                "role": u.role,
                            }
                            _persist_session_cookie(u.id)
                            st.rerun()
                except Exception as e:
                    show_busy_error(e, context="Login")

        with st.expander("Resend verification link"):
            st.caption("If you created a workspace but did not finish verification, request a new link.")
            with st.form("resend_verification"):
                r_email = st.text_input("Your email", key="resend_v_email")
                if st.form_submit_button("Issue new verification link"):
                    ok_r, out = _issue_new_verification_token(email=r_email)
                    if ok_r:
                        if out.startswith("http"):
                            st.markdown(f"[Open verification link]({out})")
                            st.code(out, language=None)
                        else:
                            st.info(out)
                    else:
                        st.error(out)

        with st.expander("Verify with token (manual)"):
            st.caption("Paste the long token from your saved link, or type ?verify_token=… value.")
            with st.form("manual_verify_token"):
                raw_t = st.text_input("Verification token", type="password", key="manual_vtok")
                if st.form_submit_button("Mark email as verified"):
                    if _consume_verification_token(raw_t.strip()):
                        st.success("Email verified. You can sign in above.")
                    else:
                        st.error("Invalid or expired token.")

    with tab_signup:
        if not signup_ok:
            st.info(
                "Self-service signup is turned off (`SELF_SIGNUP_ENABLED=false`). "
                "Ask an administrator for an account or use bootstrap secrets."
            )
        else:
            st.caption("Start a new workspace. You will be the organization admin.")
            with st.form("signup_workspace"):
                ws_slug = st.text_input(
                    "Workspace URL",
                    placeholder="acme-corp",
                    help="Short id: letters, numbers, hyphens. Must be unique.",
                    key="su_slug",
                )
                ws_name = st.text_input("Organization name", placeholder="Acme Corp", key="su_org")
                su_email = st.text_input("Your email", key="su_email")
                su_pwd = st.text_input("Password (min 8 characters)", type="password", key="su_pwd")
                su_pwd2 = st.text_input("Confirm password", type="password", key="su_pwd2")
                if st.form_submit_button("Create workspace"):
                    if su_pwd != su_pwd2:
                        st.error("Passwords do not match.")
                    else:
                        ok, msg, extra = register_workspace(
                            slug_raw=ws_slug,
                            display_name=ws_name,
                            email=su_email,
                            password=su_pwd,
                        )
                        if ok:
                            st.success(msg)
                            if extra:
                                url = (extra.get("verification_url") or "").strip()
                                if url:
                                    st.markdown(f"[Open verification link]({url})")
                                    st.code(url, language=None)
                                else:
                                    st.warning(
                                        "Set **`APP_BASE_URL`** in Streamlit secrets (your app URL, no trailing slash) "
                                        "to generate a full link. Until then, append this to your browser:"
                                    )
                                    tok = extra.get("verification_token") or ""
                                    st.code(f"?verify_token={tok}", language=None)
                                st.caption(
                                    f"Link expires in ~{extra.get('expires_hours', 48)} hours. "
                                    "Optional: send yourself this URL via email from your own mailbox."
                                )
                        else:
                            st.error(msg)
    st.stop()


def _optional_access_gate(brand: dict[str, object]) -> None:
    """If APP_ACCESS_PASSWORD is set, require it before rendering the rest of the app."""
    expected = _secret_or_env("APP_ACCESS_PASSWORD")
    if not expected:
        return
    if st.session_state.get("_access_granted"):
        return
    au = _current_user()
    if au and au.get("user_id") is not None:
        st.session_state._access_granted = True
        return
    st.title(str(brand.get("app_name") or DEFAULT_BRANDING["app_name"]))
    st.caption("This deployment is password-protected.")
    gate_btn = str(brand.get("access_gate_button") or "Continue").strip() or "Continue"
    with st.form("access_gate"):
        entered = st.text_input("Access password", type="password")
        if st.form_submit_button(gate_btn):
            if hmac.compare_digest(
                entered.encode("utf-8"),
                expected.encode("utf-8"),
            ):
                st.session_state._access_granted = True
                st.rerun()
            else:
                st.error("Incorrect password.")
    st.stop()


def _coerce_bool(value: object, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    s = str(value).strip().lower()
    if not s:
        return default
    return s in ("1", "true", "yes", "on")


def _safe_hex_color(value: object, default: str) -> str:
    if value is None:
        return default
    s = str(value).strip()
    if re.fullmatch(r"#([0-9A-Fa-f]{3}|[0-9A-Fa-f]{6})", s):
        return s
    return default


def _hex_to_rgb_tuple(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _rgba(hex_color: str, alpha: float) -> str:
    r, g, b = _hex_to_rgb_tuple(hex_color)
    return f"rgba({r}, {g}, {b}, {alpha})"


def _branding_from_secrets_file() -> dict[str, object]:
    b = _LOCAL_SECRETS.get("branding") if isinstance(_LOCAL_SECRETS, dict) else None
    return dict(b) if isinstance(b, dict) else {}


def _branding_from_streamlit_secrets() -> dict[str, object]:
    try:
        if "branding" not in st.secrets:
            return {}
        sec = st.secrets["branding"]
    except (StreamlitSecretNotFoundError, OSError, PermissionError):
        return {}
    if not isinstance(sec, Mapping):
        return {}
    raw = dict(sec)
    return {str(k): v for k, v in raw.items() if str(k) in _BRANDING_KEYS}


def _branding_env_overrides() -> dict[str, object]:
    out: dict[str, object] = {}
    mapping = {
        "BRAND_APP_NAME": "app_name",
        "BRAND_TAGLINE": "tagline",
        "BRAND_PAGE_TITLE": "page_title",
        "BRAND_PAGE_ICON": "page_icon",
        "BRAND_LOGO_URL": "logo_url",
        "BRAND_LOGO_PATH": "logo_path",
        "BRAND_PRIMARY_COLOR": "primary_color",
        "BRAND_SECONDARY_COLOR": "secondary_color",
        "BRAND_ACCENT_COLOR": "accent_color",
        "BRAND_HIDE_POWERED_BY": "hide_powered_by",
        "BRAND_ACCESS_GATE_BUTTON": "access_gate_button",
        "BRAND_SIGN_IN_BUTTON": "sign_in_button",
    }
    for env_key, brand_key in mapping.items():
        val = os.getenv(env_key)
        if val is None or val.strip() == "":
            continue
        if brand_key == "hide_powered_by":
            out[brand_key] = _coerce_bool(val)
        else:
            out[brand_key] = val.strip()
    return out


def get_initial_branding() -> dict[str, object]:
    """Branding safe to compute before the first Streamlit command.

    Used for st.set_page_config so the favicon/tab icon can be customized.
    """
    merged: dict[str, object] = dict(DEFAULT_BRANDING)
    merged.update({k: v for k, v in _branding_from_secrets_file().items() if k in _BRANDING_KEYS})
    merged.update(_branding_env_overrides())
    merged["page_icon"] = str(merged.get("page_icon") or DEFAULT_BRANDING["page_icon"])
    merged["page_title"] = str(merged.get("page_title") or DEFAULT_BRANDING["page_title"])
    return merged


_initial_brand = get_initial_branding()
st.set_page_config(
    page_title=str(_initial_brand.get("page_title")),
    page_icon=str(_initial_brand.get("page_icon")),
    layout="wide",
)


def get_branding() -> dict[str, object]:
    merged: dict[str, object] = dict(DEFAULT_BRANDING)
    merged.update({k: v for k, v in _branding_from_secrets_file().items() if k in _BRANDING_KEYS})
    merged.update({k: v for k, v in _branding_from_streamlit_secrets().items() if k in _BRANDING_KEYS})
    merged.update(_branding_env_overrides())

    merged["primary_color"] = _safe_hex_color(
        merged.get("primary_color"), str(DEFAULT_BRANDING["primary_color"])
    )
    merged["secondary_color"] = _safe_hex_color(
        merged.get("secondary_color"), str(DEFAULT_BRANDING["secondary_color"])
    )
    merged["accent_color"] = _safe_hex_color(
        merged.get("accent_color"), str(DEFAULT_BRANDING["accent_color"])
    )
    merged["hide_powered_by"] = _coerce_bool(merged.get("hide_powered_by"), False)
    merged["page_icon"] = str(merged.get("page_icon") or DEFAULT_BRANDING["page_icon"])
    return merged


def _sync_browser_tab_title(title: str) -> None:
    safe = json.dumps(title or "SOP Generator")
    components.html(
        f"<script>try{{parent.document.title = {safe};}}catch(e){{}}</script>",
        height=0,
        width=0,
    )


def build_branding_css(brand: dict[str, object]) -> str:
    pr = str(brand.get("primary_color"))
    sec = str(brand.get("secondary_color"))
    ac = str(brand.get("accent_color"))

    shadow = _rgba(sec, 0.20)
    shadow_h = _rgba(sec, 0.26)
    shadow_a = _rgba(sec, 0.18)

    return f"""
<style>
/* ---- White-label UI (brand colors) ---- */
.stApp {{
  /* Dark, muted base so the app doesn't feel overly bright */
  background: radial-gradient(900px 500px at 10% 10%, {_rgba(pr, 0.10)}, transparent 60%),
              radial-gradient(800px 520px at 90% 20%, {_rgba(sec, 0.10)}, transparent 55%),
              radial-gradient(900px 600px at 50% 90%, {_rgba(ac, 0.08)}, transparent 60%),
              linear-gradient(180deg, #0B0F1A 0%, #0A0D14 100%);
  background-size: 120% 120%;
  animation: bgShift 14s ease-in-out infinite;
}}

@keyframes bgShift {{
  0%   {{ background-position: 0% 0%; }}
  50%  {{ background-position: 100% 40%; }}
  100% {{ background-position: 0% 0%; }}
}}

section[data-testid="stSidebar"] > div {{
  background: rgba(18, 24, 39, 0.72);
  backdrop-filter: blur(10px);
  border-right: 1px solid rgba(255, 255, 255, 0.08);
}}

div.block-container {{
  padding-top: 1.25rem;
  padding-bottom: 2.5rem;
}}

div.stButton > button {{
  border: 0;
  border-radius: 14px;
  padding: 0.65rem 1rem;
  background: linear-gradient(135deg, {pr} 0%, {sec} 55%, {ac} 100%);
  color: white !important;
  box-shadow: 0 10px 24px {shadow};
  transition: transform 120ms ease, box-shadow 120ms ease, filter 120ms ease;
}}
div.stButton > button:hover {{
  transform: translateY(-1px);
  box-shadow: 0 14px 28px {shadow_h};
  filter: saturate(1.05);
}}
div.stButton > button:active {{
  transform: translateY(0px) scale(0.99);
  box-shadow: 0 8px 18px {shadow_a};
}}

div[data-baseweb="input"] input,
div[data-baseweb="textarea"] textarea,
div[data-baseweb="select"] > div {{
  border-radius: 14px !important;
}}

div[data-testid="stExpander"] {{
  border-radius: 16px;
  border: 1px solid rgba(255,255,255,0.10);
  background: rgba(18, 24, 39, 0.55);
  backdrop-filter: blur(10px);
}}

h1, h2, h3 {{
  letter-spacing: -0.02em;
}}

/* Sidebar: Sign out & secondary actions — not full-width gradient CTAs */
section[data-testid="stSidebar"] button[data-testid="baseButton-secondary"],
section[data-testid="stSidebar"] div.stButton > button[kind="secondary"] {{
  background: rgba(255, 255, 255, 0.07) !important;
  border: 1px solid rgba(255, 255, 255, 0.14) !important;
  box-shadow: none !important;
  color: rgba(230, 234, 242, 0.95) !important;
  padding: 0.4rem 0.75rem !important;
  font-size: 0.85rem !important;
}}
section[data-testid="stSidebar"] button[data-testid="baseButton-secondary"]:hover {{
  filter: brightness(1.12);
  transform: none !important;
}}
</style>
"""


SVG_CODE = """
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 240 80" width="240" height="80">
  <rect x="0" y="0" width="240" height="80" fill="#ffffff" rx="12" ry="12"/>
  <g transform="translate(20,40)">
    <circle cx="0" cy="0" r="24" fill="#0A74DA"/>
    <circle cx="8" cy="-6" r="12" fill="#ffffff"/>
  </g>
  <text x="70" y="48" font-family="Arial" font-size="36" font-weight="600" fill="#222222">SOP</text>
  <text x="70" y="70" font-family="Arial" font-size="14" fill="#555555">AI Generator</text>
</svg>
""".strip()


def render_svg_data_uri(svg: str) -> str:
    b64 = base64.b64encode(svg.encode("utf-8")).decode("utf-8")
    return f"data:image/svg+xml;base64,{b64}"


def _resolve_local_logo_path(path_str: str) -> Path | None:
    """Resolve logo_path to an existing file (repo root, .streamlit/, then cwd)."""
    raw = (path_str or "").strip()
    if not raw:
        return None
    repo_root = Path(__file__).resolve().parent
    p_in = Path(raw)
    candidates: list[Path] = []
    if p_in.is_absolute():
        candidates.append(p_in)
    else:
        candidates.extend(
            [
                repo_root / raw,
                repo_root / ".streamlit" / raw,
                Path(os.getcwd()) / raw,
                Path(os.getcwd()) / ".streamlit" / raw,
            ]
        )
    for c in candidates:
        try:
            if c.is_file():
                return c
        except OSError:
            continue
    return None


def resolve_brand_logo_url(brand: dict[str, object]) -> str:
    url = str(brand.get("logo_url") or "").strip()
    if url:
        return url
    path_str = str(brand.get("logo_path") or "").strip()
    resolved = _resolve_local_logo_path(path_str)
    if resolved is not None:
        with open(resolved, "rb") as f:
            raw = f.read()
        ext = resolved.suffix.lower()
        mime = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".webp": "image/webp",
            ".svg": "image/svg+xml",
        }.get(ext, "application/octet-stream")
        b64 = base64.b64encode(raw).decode("ascii")
        return f"data:{mime};base64,{b64}"
    return render_svg_data_uri(SVG_CODE)


def header_tagline(brand: dict[str, object]) -> str | None:
    tag = str(brand.get("tagline") or "").strip()
    if _coerce_bool(brand.get("hide_powered_by"), False):
        return tag or None
    if tag:
        return tag
    return "Powered by Groq"


_brand = get_branding()
_tab_title = str(_brand.get("page_title") or _brand.get("app_name") or DEFAULT_BRANDING["page_title"])
_sync_browser_tab_title(_tab_title)
st.markdown(build_branding_css(_brand), unsafe_allow_html=True)

# Restore login before deploy gate so refresh keeps session + skips APP_ACCESS_PASSWORD when appropriate.
if _login_required():
    _init_db()
    _ensure_bootstrap_admin()
    _restore_auth_from_signed_cookie()

_optional_access_gate(_brand)

# SaaS auth gate (preferred over shared password for client deployments)
_require_auth_ui(_brand)

# Convenience globals for downstream actions (generation, feedback, etc.)
_auth = _current_user() or {}
TENANT_ID: int | None = _auth.get("tenant_id")
USER_ID: int | None = _auth.get("user_id")
USER_ROLE: str = str(_auth.get("role") or "member")

def _normalize_pdf_text(s: str) -> str:
    """Best-effort Latin-1 for core PDF fonts (Helvetica)."""
    if not s:
        return ""
    t = unicodedata.normalize("NFKC", str(s))
    for a, b in (
        ("\u2013", "-"),
        ("\u2014", "-"),
        ("\u2212", "-"),
        ("\u2018", "'"),
        ("\u2019", "'"),
        ("\u201c", '"'),
        ("\u201d", '"'),
        ("\u2022", "-"),
        ("\u00b7", "-"),
        ("\u2026", "..."),
        ("\u00a0", " "),
    ):
        t = t.replace(a, b)
    return t.encode("latin-1", "replace").decode("latin-1")


def _extract_title_from_sop_markdown(md: str) -> str:
    lines = (md or "").splitlines()
    in_title = False
    for ln in lines:
        s = ln.strip()
        if s.startswith("## Title"):
            in_title = True
            continue
        if in_title:
            if s.startswith("## ") and not s.startswith("## Title"):
                break
            if s:
                return _normalize_pdf_text(s[:240])
    return "Standard Operating Procedure"


def _iter_sop_pdf_body_lines(md: str):
    """Skip the ## Title section (title is rendered on the cover)."""
    skipping_title = False
    for raw in (md or "").splitlines():
        st = raw.strip()
        if st.startswith("## Title"):
            skipping_title = True
            continue
        if skipping_title:
            if st.startswith("## ") and not st.startswith("## Title"):
                skipping_title = False
            else:
                continue
        yield raw


class SopReportPDF(FPDF):
    """A4 SOP layout: cover block on page 1, section hierarchy, header/footer from page 2."""

    def __init__(self, *, organization: str):
        super().__init__(unit="mm", format="A4")
        self.organization = organization
        self.set_auto_page_break(auto=True, margin=18)
        self.set_margins(18, 22, 18)

    def header(self):
        if self.page_no() == 1:
            return
        self.set_font("Helvetica", "", 8)
        self.set_text_color(115, 118, 128)
        self.cell(0, 5, self.organization, align="R")
        self.ln(1)
        self.set_draw_color(220, 222, 228)
        self.set_line_width(0.25)
        y = self.get_y()
        self.line(self.l_margin, y, self.w - self.r_margin, y)
        self.ln(4)

    def footer(self):
        self.set_y(-12)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(140, 142, 150)
        self.cell(0, 8, f"Page {self.page_no()} of {{nb}}", align="C")


def _render_sop_to_pdf(pdf: SopReportPDF, md: str) -> None:
    pdf.add_page()
    title = _extract_title_from_sop_markdown(md)
    content_w = pdf.w - pdf.l_margin - pdf.r_margin

    # --- Cover (page 1) ---
    pdf.set_y(20)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(95, 99, 108)
    pdf.cell(0, 5, pdf.organization, align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(22, 27, 40)
    pdf.multi_cell(0, 9, title, align="C")
    pdf.ln(3)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(110, 114, 122)
    pdf.cell(0, 6, "Standard Operating Procedure", align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 8)
    # ASCII-only — locale-aware "%B" month names break Helvetica / latin-1 on non-English hosts.
    stamp = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    pdf.cell(0, 5, stamp, align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(10)
    pdf.set_draw_color(195, 198, 206)
    pdf.set_line_width(0.45)
    y_rule = pdf.get_y()
    pdf.line(pdf.l_margin, y_rule, pdf.w - pdf.r_margin, y_rule)
    pdf.ln(11)

    pdf.set_text_color(38, 40, 48)
    body_font = 11
    line_h = 5.8

    for raw in _iter_sop_pdf_body_lines(md):
        line = raw.rstrip()
        s = line.strip()
        if not s:
            pdf.ln(3)
            continue
        if s.startswith("### "):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(18, 58, 92)
            pdf.multi_cell(0, 7, _normalize_pdf_text(s[4:]))
            pdf.set_font("Helvetica", "", body_font)
            pdf.set_text_color(38, 40, 48)
            pdf.ln(1)
            continue
        if s.startswith("## "):
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(22, 33, 62)
            pdf.multi_cell(0, 8, _normalize_pdf_text(s[3:]))
            pdf.set_font("Helvetica", "", body_font)
            pdf.set_text_color(38, 40, 48)
            pdf.ln(3)
            continue
        if s.startswith("# "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(22, 27, 40)
            pdf.multi_cell(0, 8, _normalize_pdf_text(s[2:]))
            pdf.set_font("Helvetica", "", body_font)
            pdf.set_text_color(38, 40, 48)
            pdf.ln(3)
            continue
        if s.startswith("- ") or s.startswith("* "):
            pdf.set_font("Helvetica", "", body_font)
            pdf.set_text_color(38, 40, 48)
            indent = 5.5
            pdf.set_x(pdf.l_margin + indent)
            bullet = "- " + _normalize_pdf_text(s[2:])
            pdf.multi_cell(content_w - indent, line_h, bullet)
            pdf.set_x(pdf.l_margin)
            continue

        pdf.set_font("Helvetica", "", body_font)
        pdf.set_text_color(38, 40, 48)
        pdf.multi_cell(0, line_h, _normalize_pdf_text(line))


def _create_pdf_bytes_simple_fallback(text: str) -> bytes:
    """Plain layout if the styled SOP PDF fails (keeps download buttons working)."""
    pdf = FPDF(unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(18, 18, 18)
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    clean = _normalize_pdf_text(text or "")
    pdf.multi_cell(0, 6, txt=clean)
    out = pdf.output(dest="S")
    if isinstance(out, str):
        return out.encode("latin-1")
    return bytes(out)


def create_pdf_bytes(text):
    try:
        brand = get_branding()
        organization = _normalize_pdf_text(str(brand.get("app_name") or DEFAULT_BRANDING["app_name"]))
        pdf = SopReportPDF(organization=organization)
        pdf.alias_nb_pages()
        doc_title = _extract_title_from_sop_markdown(text or "")
        pdf.set_title(doc_title[:180])
        pdf.set_creator(organization[:120])
        _render_sop_to_pdf(pdf, text or "")

        out = pdf.output(dest="S")
        if isinstance(out, str):
            return out.encode("latin-1")
        return bytes(out)
    except Exception as ex:
        log_exception(ex, context="create_pdf_bytes styled layout")
        try:
            return _create_pdf_bytes_simple_fallback(text or "")
        except Exception as ex2:
            log_exception(ex2, context="create_pdf_bytes fallback")
            return b""


def create_docx_bytes(title: str, text: str) -> bytes:
    doc = Document()
    if title.strip():
        doc.add_heading(title.strip(), level=1)

    for raw_line in (text or "").splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        else:
            doc.add_paragraph(line)

    buff = BytesIO()
    doc.save(buff)
    return buff.getvalue()


def sop_fingerprint(text: str) -> str:
    return hashlib.sha256((text or "").encode("utf-8")).hexdigest()


REQUIRED_SOP_HEADERS = ["## Title", "## Purpose", "## Roles", "## Procedures"]


def is_valid_sop_markdown(md: str) -> bool:
    """Validate SOP uses exactly the required H2 headers, in order."""
    text = (md or "").strip()
    if not text:
        return False
    # Must include required headers in order.
    pos = -1
    for h in REQUIRED_SOP_HEADERS:
        nxt = text.find(h)
        if nxt < 0 or nxt <= pos:
            return False
        pos = nxt
    # Must not include any other H2 headers.
    h2_lines = [ln.strip() for ln in text.splitlines() if ln.strip().startswith("## ")]
    return h2_lines == REQUIRED_SOP_HEADERS


def _chunk_text(text: str, *, chunk_chars: int = 900, overlap_chars: int = 150) -> list[str]:
    t = (text or "").strip()
    if not t:
        return []
    chunks: list[str] = []
    i = 0
    while i < len(t):
        end = min(len(t), i + chunk_chars)
        chunk = t[i:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(t):
            break
        i = max(0, end - overlap_chars)
    return chunks


@st.cache_data(show_spinner=False, ttl=3600, max_entries=128)
def extract_pdf_text_cached(*, file_sha256: str, pdf_bytes: bytes) -> str:
    # file_sha256 is part of the cache key; pdf_bytes is the payload.
    reader = PdfReader(BytesIO(pdf_bytes))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n".join(parts).strip()


@st.cache_data(show_spinner=False, ttl=3600, max_entries=64)
def build_rag_index_cached(*, corpus_sha256: str, chunks: list[str]) -> dict:
    # Store vectorizer vocabulary + matrix; st.cache_data will pickle it.
    vectorizer = TfidfVectorizer(stop_words="english", max_features=5000)
    matrix = vectorizer.fit_transform(chunks) if chunks else None
    return {"vectorizer": vectorizer, "matrix": matrix}


def retrieve_company_snippets(
    *,
    query: str,
    docs: list[dict],
    top_k: int = 6,
) -> list[dict]:
    # docs: [{name, chunks:[...]}]
    all_chunks: list[str] = []
    meta: list[dict] = []
    for d in docs:
        name = d.get("name", "manual.pdf")
        for idx, ch in enumerate(d.get("chunks", []) or []):
            all_chunks.append(ch)
            meta.append({"doc": name, "chunk_index": idx, "text": ch})

    if not all_chunks or not query.strip():
        return []

    corpus_sha = hashlib.sha256(("\n".join(all_chunks)).encode("utf-8")).hexdigest()
    index = build_rag_index_cached(corpus_sha256=corpus_sha, chunks=all_chunks)
    vectorizer: TfidfVectorizer = index["vectorizer"]
    matrix = index["matrix"]
    if matrix is None:
        return []

    qv = vectorizer.transform([query])
    sims = cosine_similarity(qv, matrix)[0]
    ranked = sorted(range(len(sims)), key=lambda i: sims[i], reverse=True)[:top_k]

    results: list[dict] = []
    for i in ranked:
        if sims[i] <= 0:
            continue
        m = meta[i]
        results.append(
            {
                "doc": m["doc"],
                "chunk_index": m["chunk_index"],
                "score": float(sims[i]),
                "text": m["text"],
            }
        )
    return results



def get_groq_api_key() -> str | None:
    """Resolve Groq API key from Streamlit secrets.toml (or env injected by Streamlit), then bare env."""
    try:
        if "GROQ_API_KEY" in st.secrets:
            k = _normalize_secret_value(st.secrets["GROQ_API_KEY"])
            if k:
                return k
        if "groq" in st.secrets:
            section = st.secrets["groq"]
            if isinstance(section, Mapping) and "api_key" in section:
                k = _normalize_secret_value(section["api_key"])
                if k:
                    return k
    except (StreamlitSecretNotFoundError, OSError, PermissionError):
        pass

    return _normalize_secret_value(os.getenv("GROQ_API_KEY"))


APP_LOG_PATH = os.path.join(".streamlit", "app.log")


def load_company_profile(*, tenant_id: int | None) -> dict:
    if tenant_id is None:
        return {}
    try:
        _init_db()
        with _db() as s:
            row = s.execute(select(TenantSettings).where(TenantSettings.tenant_id == tenant_id)).scalar_one_or_none()
            if row is None:
                row = TenantSettings(tenant_id=tenant_id)
                s.add(row)
                s.commit()
            # Ensure quotas row exists too (non-blocking)
            q = s.execute(select(TenantQuota).where(TenantQuota.tenant_id == tenant_id)).scalar_one_or_none()
            if q is None:
                s.add(
                    TenantQuota(
                        tenant_id=tenant_id,
                        generations_per_day=_coerce_int(_secret_or_env("DEFAULT_GENERATIONS_PER_DAY"), 50),
                        transcriptions_per_day=_coerce_int(_secret_or_env("DEFAULT_TRANSCRIPTIONS_PER_DAY"), 50),
                        vision_analyses_per_day=_coerce_int(_secret_or_env("DEFAULT_VISION_PER_DAY"), 50),
                    )
                )
                s.commit()
            return {
                "audience": row.audience or "",
                "tools_used": row.tools_used or "",
                "compliance_standard": row.compliance_standard or "",
                "tone": row.tone or "Professional",
            }
    except Exception as e:
        log_exception(e, context="Load tenant settings")
        return {}


def save_company_profile(*, tenant_id: int | None, profile: dict) -> None:
    if tenant_id is None:
        return
    _init_db()
    with _db() as s:
        row = s.execute(select(TenantSettings).where(TenantSettings.tenant_id == tenant_id)).scalar_one_or_none()
        if row is None:
            row = TenantSettings(tenant_id=tenant_id)
            s.add(row)
            s.flush()
        row.audience = str(profile.get("audience", "") or "")
        row.tools_used = str(profile.get("tools_used", "") or "")
        row.compliance_standard = str(profile.get("compliance_standard", "") or "")
        row.tone = str(profile.get("tone", "Professional") or "Professional")
        row.updated_at = datetime.now(timezone.utc)
        s.commit()


def append_feedback(*, tenant_id: int | None, entry: dict) -> None:
    if tenant_id is None:
        return
    try:
        _init_db()
        temp100 = int(round(float(entry.get("temperature", 0.0)) * 100))
        with _db() as s:
            s.add(
                Feedback(
                    tenant_id=tenant_id,
                    ts=datetime.fromisoformat(entry.get("ts")) if entry.get("ts") else datetime.now(timezone.utc),
                    rating=str(entry.get("rating") or ""),
                    reason=str(entry.get("reason") or ""),
                    template_name=str(entry.get("template_name") or ""),
                    strictness=str(entry.get("strictness") or ""),
                    tone=str(entry.get("tone") or ""),
                    compliance_standard=str(entry.get("compliance_standard") or ""),
                    audience=str(entry.get("audience") or ""),
                    tools_used=str(entry.get("tools_used") or ""),
                    include_definitions=bool(entry.get("include_definitions")),
                    include_safety_compliance=bool(entry.get("include_safety_compliance")),
                    include_records=bool(entry.get("include_records")),
                    include_checklist=bool(entry.get("include_checklist")),
                    model=str(entry.get("model") or ""),
                    temperature=temp100,
                    notes_chars=_coerce_int(entry.get("notes_chars"), 0),
                    sop_sha256=str(entry.get("sop_sha256") or ""),
                    source=str(entry.get("source") or ""),
                    meta={k: v for k, v in entry.items() if k not in {"ts", "rating", "reason"}},
                )
            )
            s.commit()
    except Exception as e:
        log_exception(e, context="Append feedback")


def load_recent_feedback(*, tenant_id: int | None, limit: int = 50) -> list[dict]:
    if tenant_id is None:
        return []
    try:
        _init_db()
        with _db() as s:
            rows = (
                s.execute(
                    select(Feedback)
                    .where(Feedback.tenant_id == tenant_id)
                    .order_by(Feedback.ts.desc())
                    .limit(int(limit))
                )
                .scalars()
                .all()
            )
        out: list[dict] = []
        for r in rows[::-1]:
            out.append(
                {
                    "ts": r.ts.isoformat(),
                    "rating": r.rating,
                    "reason": r.reason,
                    "template_name": r.template_name,
                    "strictness": r.strictness,
                    "tone": r.tone,
                    "compliance_standard": r.compliance_standard,
                    "audience": r.audience,
                    "tools_used": r.tools_used,
                    "include_definitions": r.include_definitions,
                    "include_safety_compliance": r.include_safety_compliance,
                    "include_records": r.include_records,
                    "include_checklist": r.include_checklist,
                    "model": r.model,
                    "temperature": (r.temperature or 0) / 100.0,
                    "notes_chars": r.notes_chars,
                    "sop_sha256": r.sop_sha256,
                    "source": r.source,
                    **(r.meta or {}),
                }
            )
        return out
    except Exception as e:
        log_exception(e, context="Load feedback")
        return []


def _user_sees_all_tenant_sops(role: str) -> bool:
    """Workspace admins/reviewers see all SOPs; members only see documents they created."""
    return str(role or "").strip().lower() in {"admin", "reviewer"}


def _sop_doc_owner_clause(role: str, user_id: int | None):
    """Employees (members): SopDoc.created_by_user_id matches user — same as JOIN users WHERE email = :current_email."""
    if _user_sees_all_tenant_sops(role):
        return None
    if user_id is None:
        return sql_false()
    return SopDoc.created_by_user_id == int(user_id)


def load_history(*, tenant_id: int | None, user_id: int | None, role: str, limit: int = 5) -> list[dict]:
    if tenant_id is None:
        return []
    try:
        _init_db()
        own = _sop_doc_owner_clause(role, user_id)
        with _db() as s:
            stmt = (
                select(SopVersion)
                .join(SopDoc, SopDoc.id == SopVersion.sop_doc_id)
                .where(SopVersion.tenant_id == tenant_id)
            )
            if own is not None:
                stmt = stmt.where(own)
            rows = (
                s.execute(stmt.order_by(SopVersion.created_at.desc()).limit(int(limit)))
                .scalars()
                .all()
            )
        out: list[dict] = []
        for r in rows:
            out.append(
                {
                    "id": r.id,
                    "ts": r.created_at.isoformat(),
                    "label": r.label,
                    "template_name": "",
                    "source": r.source,
                    "sop_text": r.sop_text,
                    "sop_sha256": r.sop_sha256,
                    "status": r.status,
                    "sop_doc_id": r.sop_doc_id,
                    "version": r.version,
                }
            )
        return out
    except Exception as e:
        log_exception(e, context="Load history")
        return []


def load_doc_versions(
    *, tenant_id: int | None, sop_doc_id: int, user_id: int | None, role: str, limit: int = 25
) -> list[dict]:
    if tenant_id is None:
        return []
    _init_db()
    own = _sop_doc_owner_clause(role, user_id)
    with _db() as s:
        stmt = (
            select(SopVersion)
            .join(SopDoc, SopDoc.id == SopVersion.sop_doc_id)
            .where(
                SopVersion.tenant_id == tenant_id,
                SopVersion.sop_doc_id == int(sop_doc_id),
                SopDoc.tenant_id == tenant_id,
            )
        )
        if own is not None:
            stmt = stmt.where(own)
        rows = (
            s.execute(stmt.order_by(SopVersion.version.desc()).limit(int(limit)))
            .scalars()
            .all()
        )
    out: list[dict] = []
    for r in rows:
        out.append(
            {
                "id": r.id,
                "version": r.version,
                "status": r.status,
                "ts": r.created_at.isoformat(),
                "source": r.source,
                "change_note": r.change_note,
                "sop_text": r.sop_text,
            }
        )
    return out


def diff_text(a: str, b: str) -> str:
    a_lines = (a or "").splitlines(keepends=True)
    b_lines = (b or "").splitlines(keepends=True)
    return "".join(difflib.unified_diff(a_lines, b_lines, fromfile="version_a", tofile="version_b"))


def list_sop_docs(
    *, tenant_id: int | None, user_id: int | None, role: str, query: str = "", limit: int = 100
) -> list[dict]:
    if tenant_id is None:
        return []
    q = (query or "").strip().lower()
    own = _sop_doc_owner_clause(role, user_id)
    _init_db()
    with _db() as s:
        stmt = select(SopDoc).where(SopDoc.tenant_id == tenant_id)
        if own is not None:
            stmt = stmt.where(own)
        docs = (
            s.execute(stmt.order_by(SopDoc.created_at.desc()).limit(int(limit)))
            .scalars()
            .all()
        )
        out: list[dict] = []
        for d in docs:
            if q and q not in (d.title or "").lower() and q not in (d.template_name or "").lower():
                continue
            latest = (
                s.execute(
                    select(SopVersion)
                    .where(SopVersion.tenant_id == tenant_id, SopVersion.sop_doc_id == d.id)
                    .order_by(SopVersion.version.desc())
                    .limit(1)
                )
                .scalars()
                .one_or_none()
            )
            out.append(
                {
                    "id": d.id,
                    "title": d.title,
                    "template_name": d.template_name,
                    "created_at": d.created_at.isoformat() if d.created_at else "",
                    "latest_version": int(getattr(latest, "version", 0) or 0),
                    "latest_status": str(getattr(latest, "status", "") or ""),
                    "latest_ts": getattr(latest, "created_at", None).isoformat() if getattr(latest, "created_at", None) else "",
                    "latest_text": str(getattr(latest, "sop_text", "") or ""),
                }
            )
        return out

def _today_key() -> str:
    return date.today().isoformat()


def get_or_create_quota(*, tenant_id: int | None) -> dict:
    if tenant_id is None:
        return {
            "generations_per_day": 0,
            "transcriptions_per_day": 0,
            "vision_analyses_per_day": 0,
        }
    _init_db()
    with _db() as s:
        q = s.execute(select(TenantQuota).where(TenantQuota.tenant_id == tenant_id)).scalar_one_or_none()
        if q is None:
            q = TenantQuota(
                tenant_id=tenant_id,
                generations_per_day=_coerce_int(_secret_or_env("DEFAULT_GENERATIONS_PER_DAY"), 50),
                transcriptions_per_day=_coerce_int(_secret_or_env("DEFAULT_TRANSCRIPTIONS_PER_DAY"), 50),
                vision_analyses_per_day=_coerce_int(_secret_or_env("DEFAULT_VISION_PER_DAY"), 50),
            )
            s.add(q)
            s.commit()
        return {
            "generations_per_day": int(q.generations_per_day or 0),
            "transcriptions_per_day": int(q.transcriptions_per_day or 0),
            "vision_analyses_per_day": int(q.vision_analyses_per_day or 0),
        }


def set_quota(
    *,
    tenant_id: int | None,
    generations_per_day: int,
    transcriptions_per_day: int,
    vision_analyses_per_day: int,
) -> None:
    if tenant_id is None:
        return
    _init_db()
    with _db() as s:
        q = s.execute(select(TenantQuota).where(TenantQuota.tenant_id == tenant_id)).scalar_one_or_none()
        if q is None:
            q = TenantQuota(tenant_id=tenant_id)
            s.add(q)
            s.flush()
        q.generations_per_day = max(0, int(generations_per_day))
        q.transcriptions_per_day = max(0, int(transcriptions_per_day))
        q.vision_analyses_per_day = max(0, int(vision_analyses_per_day))
        q.updated_at = datetime.now(timezone.utc)
        s.commit()


def usage_counts_today(*, tenant_id: int | None) -> dict[str, int]:
    if tenant_id is None:
        return {"generate": 0, "transcribe": 0, "vision": 0}
    _init_db()
    day = _today_key()
    with _db() as s:
        rows = (
            s.execute(select(DailyUsage).where(DailyUsage.tenant_id == tenant_id, DailyUsage.day == day))
            .scalars()
            .all()
        )
    out = {"generate": 0, "transcribe": 0, "vision": 0}
    for r in rows:
        out[str(r.action)] = int(r.count or 0)
    return out


def check_and_consume_quota(*, tenant_id: int | None, action: str, amount: int = 1) -> tuple[bool, str]:
    """Returns (ok, message). If ok=True usage is incremented."""
    if tenant_id is None:
        return False, "Missing tenant."
    action = (action or "").strip().lower()
    if action not in {"generate", "transcribe", "vision"}:
        return False, "Invalid action."
    amount = max(1, int(amount or 1))

    quota = get_or_create_quota(tenant_id=tenant_id)
    limit = {
        "generate": quota["generations_per_day"],
        "transcribe": quota["transcriptions_per_day"],
        "vision": quota["vision_analyses_per_day"],
    }[action]
    if limit <= 0:
        return False, "This feature is disabled for your workspace."

    _init_db()
    day = _today_key()
    with _db() as s:
        row = s.execute(
            select(DailyUsage).where(DailyUsage.tenant_id == tenant_id, DailyUsage.day == day, DailyUsage.action == action)
        ).scalar_one_or_none()
        if row is None:
            row = DailyUsage(tenant_id=tenant_id, day=day, action=action, count=0)
            s.add(row)
            s.flush()
        current = int(row.count or 0)
        if current + amount > int(limit):
            remaining = max(0, int(limit) - current)
            return False, f"Daily quota reached for {action}. Remaining today: {remaining}."
        row.count = current + amount
        row.updated_at = datetime.now(timezone.utc)
        s.commit()
        remaining = max(0, int(limit) - int(row.count or 0))
        return True, f"Quota ok. Remaining today: {remaining}."


def save_history(items: list[dict]) -> None:
    # No-op: history is DB-backed now.
    return


def add_to_history(*, tenant_id: int | None, user_id: int | None, role: str = "member", entry: dict) -> None:
    """Persist an SOP version into the tenant workspace (proper versioning per SOP doc)."""
    if tenant_id is None:
        return
    _init_db()
    label = str(entry.get("label") or "SOP").strip() or "SOP"
    template_name = str(entry.get("template_name") or "IT SOP")
    sop_text = str(entry.get("sop_text") or "")
    sop_sha = str(entry.get("sop_sha256") or sop_fingerprint(sop_text))
    source = str(entry.get("source") or "generated")
    change_note = str(entry.get("change_note") or "")
    sop_doc_id = entry.get("sop_doc_id")
    own = _sop_doc_owner_clause(role, user_id)

    with _db() as s:
        doc: SopDoc | None = None
        if sop_doc_id is not None:
            try:
                q = select(SopDoc).where(SopDoc.tenant_id == tenant_id, SopDoc.id == int(sop_doc_id))
                if own is not None:
                    q = q.where(own)
                doc = s.execute(q).scalar_one_or_none()
            except Exception:
                doc = None

        if doc is None:
            # Reuse doc by (tenant, title, template) when possible — scoped to this user for members.
            q = select(SopDoc).where(
                SopDoc.tenant_id == tenant_id,
                SopDoc.title == label,
                SopDoc.template_name == template_name,
            )
            if own is not None:
                q = q.where(own)
            doc = s.execute(q).scalar_one_or_none()

        if doc is None:
            doc = SopDoc(tenant_id=tenant_id, title=label, template_name=template_name, created_by_user_id=user_id)
            s.add(doc)
            s.flush()

        next_version = (
            s.execute(
                select(SopVersion.version)
                .where(SopVersion.tenant_id == tenant_id, SopVersion.sop_doc_id == doc.id)
                .order_by(SopVersion.version.desc())
                .limit(1)
            ).scalar_one_or_none()
            or 0
        )
        next_version = int(next_version) + 1
        v = SopVersion(
            tenant_id=tenant_id,
            sop_doc_id=doc.id,
            version=next_version,
            status="draft",
            label=label,
            source=source,
            sop_text=sop_text,
            sop_sha256=sop_sha,
            change_note=change_note,
            created_at=datetime.fromisoformat(entry.get("ts")) if entry.get("ts") else datetime.now(timezone.utc),
            created_by_user_id=user_id,
        )
        s.add(v)
        s.commit()

        # Store current doc id in session for "v2/v3" behavior.
        try:
            st.session_state.current_sop_doc_id = doc.id
        except Exception:
            pass


def save_current_edits_as_new_version(
    *,
    tenant_id: int | None,
    user_id: int | None,
    role: str,
    template_name: str,
    label: str,
    sop_text: str,
    source: str,
    change_note: str,
) -> None:
    add_to_history(
        tenant_id=tenant_id,
        user_id=user_id,
        role=role,
        entry={
            "ts": datetime.now(timezone.utc).isoformat(),
            "label": label,
            "template_name": template_name,
            "source": source,
            "sop_text": sop_text,
            "sop_sha256": sop_fingerprint(sop_text),
            "change_note": change_note,
            "sop_doc_id": st.session_state.get("current_sop_doc_id"),
        },
    )


def update_sop_version_status(
    *,
    tenant_id: int | None,
    user_id: int | None,
    role: str,
    sop_version_id: int,
    new_status: str,
    approved_by_user_id: int | None,
) -> None:
    if tenant_id is None:
        return
    new_status = (new_status or "").strip().lower()
    if new_status not in {"draft", "in_review", "approved"}:
        return
    own = _sop_doc_owner_clause(role, user_id)
    _init_db()
    with _db() as s:
        stmt = (
            select(SopVersion)
            .join(SopDoc, SopDoc.id == SopVersion.sop_doc_id)
            .where(
                SopVersion.tenant_id == tenant_id,
                SopVersion.id == int(sop_version_id),
                SopDoc.tenant_id == tenant_id,
            )
        )
        if own is not None:
            stmt = stmt.where(own)
        v = s.execute(stmt).scalars().one_or_none()
        if v is None:
            return
        v.status = new_status
        if new_status == "approved":
            v.approved_at = datetime.now(timezone.utc)
            v.approved_by_user_id = approved_by_user_id
        s.commit()


def delete_sop_version(*, tenant_id: int | None, user_id: int | None, role: str, sop_version_id: int) -> None:
    if tenant_id is None:
        return
    own = _sop_doc_owner_clause(role, user_id)
    _init_db()
    with _db() as s:
        stmt = (
            select(SopVersion)
            .join(SopDoc, SopDoc.id == SopVersion.sop_doc_id)
            .where(
                SopVersion.tenant_id == tenant_id,
                SopVersion.id == int(sop_version_id),
                SopDoc.tenant_id == tenant_id,
            )
        )
        if own is not None:
            stmt = stmt.where(own)
        v = s.execute(stmt).scalars().one_or_none()
        if v is None:
            return
        s.delete(v)
        s.commit()


def list_users(*, tenant_id: int | None) -> list[dict]:
    if tenant_id is None:
        return []
    _init_db()
    with _db() as s:
        rows = (
            s.execute(select(User).where(User.tenant_id == tenant_id).order_by(User.created_at.desc()))
            .scalars()
            .all()
        )
    out: list[dict] = []
    for u in rows:
        out.append(
            {
                "id": u.id,
                "email": u.email,
                "role": u.role,
                "is_active": bool(u.is_active),
                "created_at": u.created_at.isoformat() if u.created_at else "",
            }
        )
    return out


def create_user(
    *,
    tenant_id: int | None,
    email: str,
    password: str,
    role: str,
) -> None:
    if tenant_id is None:
        return
    email_n = _normalize_email(email)
    role_n = (role or "member").strip().lower()
    if role_n not in {"admin", "reviewer", "member"}:
        role_n = "member"
    if not email_n or not password:
        return
    _init_db()
    with _db() as s:
        existing = s.execute(select(User).where(User.tenant_id == tenant_id, User.email == email_n)).scalar_one_or_none()
        if existing is not None:
            raise ValueError("User already exists.")
        s.add(
            User(
                tenant_id=tenant_id,
                email=email_n,
                password_hash=_hash_password(password),
                role=role_n,
                is_active=True,
                email_verified=True,
            )
        )
        s.commit()


def set_user_active(*, tenant_id: int | None, user_id: int, is_active: bool) -> None:
    if tenant_id is None:
        return
    _init_db()
    with _db() as s:
        u = s.execute(select(User).where(User.tenant_id == tenant_id, User.id == int(user_id))).scalar_one_or_none()
        if u is None:
            return
        u.is_active = bool(is_active)
        s.commit()


def reset_user_password(*, tenant_id: int | None, user_id: int, new_password: str) -> None:
    if tenant_id is None:
        return
    if not new_password:
        return
    _init_db()
    with _db() as s:
        u = s.execute(select(User).where(User.tenant_id == tenant_id, User.id == int(user_id))).scalar_one_or_none()
        if u is None:
            return
        u.password_hash = _hash_password(new_password)
        s.commit()


def upsert_manual_from_pdf(
    *,
    tenant_id: int | None,
    file_name: str,
    pdf_bytes: bytes,
    chunk_chars: int = 900,
    overlap_chars: int = 150,
) -> dict:
    """Persist a PDF manual and its chunks in DB (deduped by sha256 per tenant)."""
    if tenant_id is None:
        return {"stored": False, "reason": "no-tenant"}
    sha = hashlib.sha256(pdf_bytes).hexdigest()
    text = extract_pdf_text_cached(file_sha256=sha, pdf_bytes=pdf_bytes)
    chunks = _chunk_text(text, chunk_chars=chunk_chars, overlap_chars=overlap_chars)
    _init_db()
    with _db() as s:
        existing = s.execute(
            select(ManualDoc).where(ManualDoc.tenant_id == tenant_id, ManualDoc.sha256 == sha)
        ).scalar_one_or_none()
        if existing is not None:
            return {"stored": False, "reason": "duplicate", "sha256": sha, "name": existing.name}

        doc = ManualDoc(tenant_id=tenant_id, name=str(file_name or "manual.pdf"), sha256=sha)
        s.add(doc)
        s.flush()
        for idx, ch in enumerate(chunks):
            s.add(ManualChunk(tenant_id=tenant_id, manual_doc_id=doc.id, chunk_index=idx, text=ch))
        s.commit()
        return {"stored": True, "sha256": sha, "name": doc.name, "chunks": len(chunks)}


def list_manual_docs(*, tenant_id: int | None) -> list[dict]:
    if tenant_id is None:
        return []
    _init_db()
    with _db() as s:
        docs = (
            s.execute(select(ManualDoc).where(ManualDoc.tenant_id == tenant_id).order_by(ManualDoc.created_at.desc()))
            .scalars()
            .all()
        )
        out: list[dict] = []
        for d in docs:
            out.append({"id": d.id, "name": d.name, "sha256": d.sha256, "created_at": d.created_at.isoformat()})
        return out


def load_manual_docs_with_chunks(*, tenant_id: int | None, manual_doc_ids: list[int]) -> list[dict]:
    if tenant_id is None:
        return []
    ids = [int(x) for x in (manual_doc_ids or []) if str(x).strip().isdigit()]
    if not ids:
        return []
    _init_db()
    with _db() as s:
        docs = (
            s.execute(select(ManualDoc).where(ManualDoc.tenant_id == tenant_id, ManualDoc.id.in_(ids)))
            .scalars()
            .all()
        )
        out: list[dict] = []
        for d in docs:
            chunks = (
                s.execute(
                    select(ManualChunk)
                    .where(ManualChunk.tenant_id == tenant_id, ManualChunk.manual_doc_id == d.id)
                    .order_by(ManualChunk.chunk_index.asc())
                )
                .scalars()
                .all()
            )
            out.append({"id": d.id, "name": d.name, "sha256": d.sha256, "chunks": [c.text for c in chunks]})
        return out


def delete_manual_doc(*, tenant_id: int | None, manual_doc_id: int) -> None:
    if tenant_id is None:
        return
    _init_db()
    with _db() as s:
        d = s.execute(select(ManualDoc).where(ManualDoc.tenant_id == tenant_id, ManualDoc.id == int(manual_doc_id))).scalar_one_or_none()
        if d is None:
            return
        s.delete(d)
        s.commit()


TEMPLATE_GUIDANCE: dict[str, str] = {
    "IT SOP": """
Focus on technical accuracy, security, and repeatability.
Include: Preconditions/requirements, access/permissions, tools/systems involved,
rollback plan, troubleshooting, validation steps, logging/monitoring, and SLAs/owners.
Add a short 'Change management' section (impact, approvals, maintenance window).
""".strip(),
    "HR SOP": """
Focus on compliance, privacy, fairness, and a clear human workflow.
Include: Trigger events, required forms/documents, approvals, timelines/SLAs,
confidentiality/data handling, escalation paths, templates/communications, and record retention.
Add a short 'Candidate/employee communication' checklist.
""".strip(),
    "Warehouse SOP": """
Focus on safety, efficiency, and physical process clarity.
Include: PPE/safety requirements, equipment/tools, location/bin labeling, scanning steps,
quality checks, exception handling (damages/shortages), and end-of-shift reconciliation.
Add a short 'Safety checks' section and 'Common errors to avoid'.
""".strip(),
    "Restaurant SOP": """
Focus on food safety, service consistency, and speed.
Include: food safety controls (temps, cross-contamination), prep/line setup,
service steps, cleaning schedules, allergen handling, customer escalation, and close-down tasks.
Add checklists for opening/shift/closing and 'Quality standards' (taste, plating, timing).
""".strip(),
}


def build_prompt_for_template(
    template_name: str,
    topic: str,
    notes: str,
    *,
    audience: str,
    tools_used: str,
    compliance_standard: str,
    strictness: str,
    tone: str,
    include_definitions: bool,
    include_safety_compliance: bool,
    include_records: bool,
    include_checklist: bool,
) -> str:
    template_guidance = TEMPLATE_GUIDANCE.get(template_name, "")
    tenant_id = (_current_user() or {}).get("tenant_id")
    feedback_items = load_recent_feedback(tenant_id=tenant_id, limit=60)
    recent_down_reasons = [
        (it.get("reason") or "").strip()
        for it in feedback_items
        if it.get("rating") == "down" and it.get("template_name") == template_name
    ]
    recent_down_reasons = [r for r in recent_down_reasons if r][:3]
    feedback_avoid = ""
    if recent_down_reasons:
        bullets = "\n".join([f"- {r}" for r in recent_down_reasons])
        feedback_avoid = f"""
Common issues to avoid (from user feedback on previous SOPs):
{bullets}
""".strip()

    strictness_instructions = (
        "Strictness: STRICT. Use a formal, policy-like tone. Use short, unambiguous steps. "
        "Avoid fluff. Prefer MUST/SHALL where appropriate. Include clear acceptance/verification criteria."
        if strictness == "Strict"
        else "Strictness: DETAILED. Be thorough and explanatory while staying professional. "
        "Include tips, examples, and clarifying notes where helpful."
    )

    # System prompt requirement: strict Markdown headers in a fixed order.
    # Keep this list in sync with prompt instructions below.
    sections_text = "\n".join(
        [
            "1. Title",
            "2. Purpose",
            "3. Roles",
            "4. Procedures",
        ]
    )

    notes_based_section = ""
    if len((notes or "").strip()) >= 1200:
        notes_based_section = """
Add a short section near the top titled exactly: "Based on notes:"
- 5–10 bullet points capturing the most important concrete facts from the notes (tools, roles, constraints, risks, timelines).
- Each bullet should reference the notes by quoting a short phrase in double-quotes OR by citing a specific detail (names/titles are ok; avoid secrets).
- Do NOT invent details that are not present in the notes. If something is missing, say "Not specified in notes".
""".strip()
    company_rules = st.session_state.get("company_rules_context", "").strip()
    company_rules_block = ""
    if company_rules:
        company_rules_block = f"""
Company rules (from uploaded manuals) — FOLLOW THESE STRICTLY:
{company_rules}
If any manual rule conflicts with the user's notes, call out the conflict and choose the safer/compliant path.
When a procedure step is directly based on a manual rule, add an inline citation at the end of the step like:
(Manual: <filename> chunk <N>).
""".strip()

    return f"""
Write a clear, professional Standard Operating Procedure (SOP) for the topic below.

OUTPUT FORMAT (STRICT):
- Output MUST be valid Markdown.
- Use ONLY these H2 headers, in this exact order (spelling/case must match):
  1) ## Title
  2) ## Purpose
  3) ## Roles
  4) ## Procedures
- Do NOT add any other headers (no extra ## sections).
- Under "## Procedures" use a numbered list (1., 2., 3., ...).

Target audience: {audience}
Tools/systems used: {tools_used or "Not specified"}
Compliance standard(s): {compliance_standard or "Not specified"}
Tone: {tone}
{strictness_instructions}

{feedback_avoid}

{company_rules_block}

{notes_based_section}

Required headers (include ONLY these; omit all others):
{sections_text}

Template-specific guidance:
{template_guidance}

Topic: {topic}
Notes / raw input (may be messy): {notes}
""".strip()


@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def generate_sop_cached(
    *,
    api_key: str,
    model: str,
    temperature: float,
    prompt: str,
) -> str:
    client = Groq(api_key=api_key)
    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a professional technical writer."},
            {"role": "user", "content": prompt},
        ],
        temperature=float(temperature),
    )
    return completion.choices[0].message.content or ""


@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def format_sop_to_required_markdown_cached(
    *,
    api_key: str,
    model: str,
    sop_text: str,
) -> str:
    """Reformat any SOP-ish text into the required strict Markdown structure."""
    client = Groq(api_key=api_key)
    prompt = f"""
Reformat the SOP content below into STRICT valid Markdown using ONLY these H2 headers, in this exact order:
1) ## Title
2) ## Purpose
3) ## Roles
4) ## Procedures

Rules:
- Do not add any other headers (no extra ## sections).
- Under "## Procedures" use a numbered list (1., 2., 3., ...).
- Preserve factual content; if something is missing, write "Not specified".
- Return ONLY the Markdown.

SOP content:
{sop_text}
""".strip()
    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a meticulous technical editor."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.0,
    )
    return completion.choices[0].message.content or ""


@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def review_and_fix_sop_cached(
    *,
    api_key: str,
    model: str,
    temperature: float,
    sop_text: str,
    strictness: str,
    tone: str,
    compliance_standard: str,
) -> str:
    client = Groq(api_key=api_key)
    prompt = f"""
You are reviewing an SOP for quality and completeness.

Goals:
- Find and fix gaps, unclear steps, missing roles/responsibilities, and missing records/documentation.
- Ensure steps are testable/verifyable and ordered logically.
- Ensure compliance language is appropriate for: {compliance_standard or "Not specified"} (if any).
- Keep the same overall intent, but rewrite as a corrected, improved SOP.
- If there is a "Based on notes:" section, keep it and correct it to match the SOP (do not add new facts).

Output rules:
- Output MUST be valid Markdown.
- Use ONLY these H2 headers, in this exact order (spelling/case must match):
  1) ## Title
  2) ## Purpose
  3) ## Roles
  4) ## Procedures
- Do NOT add any other headers (no extra ## sections).
- Under "## Procedures" use a numbered list (1., 2., 3., ...).
- Return ONLY the revised SOP (no analysis, no bullet list of issues).
- Use the same tone: {tone}
- Use strictness: {strictness}

SOP to review:
{sop_text}
""".strip()

    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a meticulous SOP editor and auditor."},
            {"role": "user", "content": prompt},
        ],
        temperature=float(temperature),
    )
    return completion.choices[0].message.content or ""


@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def compliance_audit_cached(
    *,
    api_key: str,
    model: str,
    temperature: float,
    sop_text: str,
    template_name: str,
    strictness: str,
    tone: str,
    compliance_standard: str,
    company_rules_context: str,
) -> str:
    client = Groq(api_key=api_key)
    prompt = f"""
You are a Compliance Auditor (AI Critic). Review the SOP below and critique it.

Context:
- Template: {template_name}
- Compliance standard(s): {compliance_standard or "Not specified"}
- Strictness: {strictness}
- Tone: {tone}

Company rules (if provided; treat as policy requirements):
{company_rules_context or "None"}

What to check:
- Compliance gaps or risky/unsafe steps
- Missing roles/responsibilities, approvals, evidence/records
- Unclear, untestable, or ambiguous steps
- Missing exceptions/edge cases and escalation paths
- Missing safety controls (PPE, cross-contamination, PHI/PII, access control, etc.) when relevant
- Conflicts with company rules (if any) and how to resolve them

Output format (use these headings exactly):
## Summary
## Findings (ranked)
- [High] ...
- [Medium] ...
- [Low] ...
## Missing items checklist
- ...
## Recommended fixes
1. ...

Return only the critique (no extra commentary).

SOP:
{sop_text}
""".strip()

    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a strict compliance auditor and SOP critic."},
            {"role": "user", "content": prompt},
        ],
        temperature=float(temperature),
    )
    return completion.choices[0].message.content or ""

@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def generate_flowchart_mermaid_cached(
    *,
    api_key: str,
    model: str,
    temperature: float,
    sop_text: str,
) -> str:
    client = Groq(api_key=api_key)
    prompt = f"""
Create a Mermaid flowchart for the SOP below.

Rules:
- Output ONLY Mermaid code.
- Start with: flowchart TD
- Keep it readable: at most ~18 nodes.
- Use decision diamonds with labels like "Yes"/"No" paths when needed.
- Include start/end nodes.
- Do NOT include markdown fences.

SOP:
{sop_text}
""".strip()

    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You convert SOPs into clear flowcharts."},
            {"role": "user", "content": prompt},
        ],
        temperature=float(temperature),
    )
    return (completion.choices[0].message.content or "").strip()


def render_mermaid(mermaid_code: str, *, height_px: int = 700) -> None:
    code = (mermaid_code or "").strip()
    if not code:
        st.info("No flowchart to display.")
        return

    # Mermaid is rendered client-side via CDN.
    html = f"""
<div class="mermaid">
{code}
</div>
<script type="module">
  import mermaid from "https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs";
  mermaid.initialize({{
    startOnLoad: true,
    theme: "default",
    flowchart: {{ curve: "basis" }}
  }});
</script>
"""
    components.html(html, height=height_px, scrolling=True)


@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def transcribe_audio_cached(
    *,
    api_key: str,
    model: str,
    file_name: str,
    file_sha256: str,
    audio_bytes: bytes,
    language: str,
) -> str:
    client = Groq(api_key=api_key)
    transcription = client.audio.transcriptions.create(
        file=(file_name, audio_bytes),
        model=model,
        response_format="json",
        language=language or None,
        temperature=0.0,
    )
    # Groq SDK returns an object with .text
    return (getattr(transcription, "text", None) or "").strip()


@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def analyze_image_to_notes_cached(
    *,
    api_key: str,
    model: str,
    file_sha256: str,
    mime_type: str,
    image_b64: str,
) -> str:
    client = Groq(api_key=api_key)
    prompt = """
You are extracting actionable SOP notes from an image.

Return concise NOTES ONLY (no preamble), as bullet points grouped by:
- What is shown
- Key entities (tools/systems/roles)
- Steps / sequence (if implied)
- Requirements / constraints
- Risks / safety / compliance signals
- Any numbers, dates, thresholds, or checklists visible

If the image is a form/table/screenshot, capture the important fields and values.
Do not invent details; if unclear, say "Unclear in image".
""".strip()

    completion = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime_type};base64,{image_b64}"},
                    },
                ],
            }
        ],
        temperature=0.1,
    )
    return (completion.choices[0].message.content or "").strip()

logo_url = resolve_brand_logo_url(_brand)

auth = _auth
tenant_id = TENANT_ID
user_id = USER_ID
user_role = str(USER_ROLE or "member")

# Load profile once per session, then use it as widget defaults.
if "company_profile_loaded" not in st.session_state:
    profile = load_company_profile(tenant_id=tenant_id)
    st.session_state.company_profile_loaded = True
    st.session_state.profile_audience = str(profile.get("audience", "") or "")
    st.session_state.profile_tools_used = str(profile.get("tools_used", "") or "")
    st.session_state.profile_compliance = str(profile.get("compliance_standard", "") or "")
    st.session_state.profile_tone = str(profile.get("tone", "Professional") or "Professional")

# Generator defaults (persist across pages via session_state keys).
if "gen_template" not in st.session_state:
    st.session_state.gen_template = "IT SOP"
if "gen_strictness" not in st.session_state:
    st.session_state.gen_strictness = "Detailed"
if "gen_audience" not in st.session_state:
    st.session_state.gen_audience = str(st.session_state.get("profile_audience") or "")
if "gen_tools_used" not in st.session_state:
    st.session_state.gen_tools_used = str(st.session_state.get("profile_tools_used") or "")
if "gen_compliance" not in st.session_state:
    pc = str(st.session_state.get("profile_compliance") or "").strip()
    st.session_state.gen_compliance = "None" if not pc else pc
if "gen_tone" not in st.session_state:
    st.session_state.gen_tone = str(st.session_state.get("profile_tone") or "Professional")
for _k, _v in (
    ("gen_inc_definitions", True),
    ("gen_inc_safety", True),
    ("gen_inc_records", True),
    ("gen_inc_checklist", True),
):
    if _k not in st.session_state:
        st.session_state[_k] = _v
if "rag_top_k" not in st.session_state:
    st.session_state.rag_top_k = 6

_existing_manual_rows = list_manual_docs(tenant_id=tenant_id)
if "manual_multiselect_ids" not in st.session_state:
    st.session_state.manual_multiselect_ids = (
        [int(m["id"]) for m in _existing_manual_rows][:3] if _existing_manual_rows else []
    )

_PAGE_LABELS = ["✨ Generator", "📂 Library", "⚙️ Settings"]
_PAGE_MAP = {"✨ Generator": "Generator", "📂 Library": "Library", "⚙️ Settings": "Settings"}

with st.sidebar:
    _has_custom_logo = bool(str(_brand.get("logo_url") or "").strip() or str(_brand.get("logo_path") or "").strip())
    if _has_custom_logo:
        st.image(logo_url, width=220)
    else:
        st.markdown(
            """
<div style="border: 1px dashed rgba(255,255,255,0.22); border-radius: 16px; padding: 18px; text-align: center; background: rgba(255,255,255,0.04); margin-bottom: 8px;">
  <div style="font-weight: 700; letter-spacing: 0.04em; opacity: 0.9;">LOGO</div>
  <div style="margin-top: 6px; font-size: 12px; opacity: 0.7;">Set <code>[branding] logo_path</code> in secrets</div>
</div>
""",
            unsafe_allow_html=True,
        )
    st.markdown(
        f"<div style='text-align:center;font-weight:700;font-size:1.05rem;'>{str(_brand.get('app_name') or DEFAULT_BRANDING['app_name'])}</div>",
        unsafe_allow_html=True,
    )
    st.caption("Professional Edition")
    st.markdown("### Workspace")
    selected_page = st.radio("Go to", _PAGE_LABELS, label_visibility="collapsed")
    active_page = _PAGE_MAP[selected_page]
    st.divider()
    with st.container(border=True):
        st.caption(f"👤 **{auth.get('email', '') or '—'}**")
        _c1, _c2 = st.columns([1, 2])
        with _c2:
            if st.button("Sign out", type="secondary", use_container_width=True, key="sidebar_sign_out"):
                st.session_state.pop("auth_user", None)
                _clear_persistent_auth_cookie()
                st.rerun()

# Settings page (widgets must run before derived generation variables are read).
if active_page == "Settings":
    st.title("⚙️ Workspace settings")
    st.caption("Company profile, knowledge base, generator defaults, and administration.")

    quota = get_or_create_quota(tenant_id=tenant_id)
    usage = usage_counts_today(tenant_id=tenant_id)
    with st.expander("Usage & quotas (today)", expanded=False):
        st.caption(f"Date: {_today_key()}")
        st.caption(f"Database: {_database_backend_label()}")
        st.write(
            {
                "generations_used": usage.get("generate", 0),
                "generations_limit": quota.get("generations_per_day", 0),
                "transcriptions_used": usage.get("transcribe", 0),
                "transcriptions_limit": quota.get("transcriptions_per_day", 0),
                "vision_used": usage.get("vision", 0),
                "vision_limit": quota.get("vision_analyses_per_day", 0),
            }
        )

    st.info(
        "1. Paste **notes** or capture **voice / vision** on **Generator**.\n"
        "2. Adjust defaults here (template, tone, manuals).\n"
        "3. Generate your SOP and export PDF/DOCX from the main workspace."
    )

    st.markdown("### Company profile")
    st.text_input(
        "Default audience",
        key="profile_audience",
        placeholder="e.g., New hires, IT admins, Shift supervisors",
    )
    st.text_input(
        "Default tools used",
        key="profile_tools_used",
        placeholder="e.g., Okta, Jira, Google Workspace, Forklifts, POS system",
    )
    st.selectbox(
        "Default compliance",
        ["", "ISO 27001", "SOC 2", "HIPAA"],
        key="profile_compliance",
        index=0,
    )
    st.selectbox(
        "Default tone",
        ["Professional", "Friendly", "Policy-like", "Concise"],
        key="profile_tone",
        index=0,
    )

    col_p1, col_p2 = st.columns(2)
    with col_p1:
        if st.button("Save profile"):
            save_company_profile(
                tenant_id=tenant_id,
                profile={
                    "audience": st.session_state.profile_audience,
                    "tools_used": st.session_state.profile_tools_used,
                    "compliance_standard": st.session_state.profile_compliance,
                    "tone": st.session_state.profile_tone,
                },
            )
            st.success("Saved.")
    with col_p2:
        if st.button("Reset profile"):
            st.session_state.profile_audience = ""
            st.session_state.profile_tools_used = ""
            st.session_state.profile_compliance = ""
            st.session_state.profile_tone = "Professional"
            save_company_profile(
                tenant_id=tenant_id,
                profile={
                    "audience": "",
                    "tools_used": "",
                    "compliance_standard": "",
                    "tone": "Professional",
                },
            )
            st.success("Reset.")

    st.markdown("### Company Brain (RAG Lite)")
    st.slider("Manual snippets to use", 2, 10, key="rag_top_k")

    existing_manuals = list_manual_docs(tenant_id=tenant_id)
    manual_label_by_id = {int(m["id"]): f'{m["name"]} ({m["sha256"][:10]}…)'.strip() for m in existing_manuals}
    st.multiselect(
        "Use these manuals",
        options=[int(m["id"]) for m in existing_manuals],
        key="manual_multiselect_ids",
        format_func=lambda mid: manual_label_by_id.get(int(mid), str(mid)),
        help="These manuals are saved per tenant and reused across sessions.",
    )

    manuals_upload = st.file_uploader(
        "Upload PDF manuals (saved to tenant library)",
        type=["pdf"],
        accept_multiple_files=True,
        help="Examples: Employee Handbook, Safety Guidelines, IT policy. These will be used as strict rules for SOPs.",
    )
    if st.button("Save uploaded manuals to library", disabled=(not manuals_upload)):
        stored = 0
        dupes = 0
        with st.spinner("Saving manuals..."):
            for f in (manuals_upload or []):
                try:
                    res = upsert_manual_from_pdf(tenant_id=tenant_id, file_name=f.name, pdf_bytes=f.getvalue())
                    if res.get("stored"):
                        stored += 1
                    elif res.get("reason") == "duplicate":
                        dupes += 1
                except Exception as e:
                    log_exception(e, context="Save manual to library")
                    continue
        st.success(f"Saved: {stored}. Duplicates skipped: {dupes}.")
        st.rerun()

    if existing_manuals:
        with st.expander("Manage manual library", expanded=False):
            m_to_delete = st.selectbox(
                "Delete a manual",
                options=[int(m["id"]) for m in existing_manuals],
                format_func=lambda mid: manual_label_by_id.get(int(mid), str(mid)),
            )
            if st.button("Delete manual permanently"):
                try:
                    delete_manual_doc(tenant_id=tenant_id, manual_doc_id=int(m_to_delete))
                    st.success("Deleted.")
                    st.rerun()
                except Exception as e:
                    show_busy_error(e, context="Delete manual")
    else:
        st.caption("No manuals in the tenant library yet.")

    st.markdown("### Generator defaults")
    st.selectbox(
        "Template",
        ["IT SOP", "HR SOP", "Warehouse SOP", "Restaurant SOP"],
        key="gen_template",
    )
    st.radio("Strictness", ["Strict", "Detailed"], key="gen_strictness", horizontal=True)
    st.text_input(
        "Audience (optional)",
        key="gen_audience",
        placeholder="e.g., New hires, IT admins, Shift supervisors",
    )
    st.text_input(
        "Tools used (optional)",
        key="gen_tools_used",
        placeholder="e.g., Okta, Jira, Google Workspace, Forklifts, POS system",
    )
    st.selectbox(
        "Compliance standard (optional)",
        ["None", "ISO 27001", "SOC 2", "HIPAA"],
        key="gen_compliance",
    )
    st.selectbox(
        "Tone",
        ["Professional", "Friendly", "Policy-like", "Concise"],
        key="gen_tone",
    )

    st.markdown("### Outline controls")
    st.checkbox("Include Definitions section", key="gen_inc_definitions")
    st.checkbox("Include Safety/Compliance section", key="gen_inc_safety")
    st.checkbox("Include Records/Documentation section", key="gen_inc_records")
    st.checkbox("Include Checklist section", key="gen_inc_checklist")

    if st.button("Clear cached results"):
        st.cache_data.clear()

    if user_role == "admin":
        with st.expander("Admin (Users)", expanded=False):
            st.caption("Create users and manage access for this tenant.")
            with st.form("create_user"):
                new_email = st.text_input("User email", placeholder="user@company.com")
                new_pwd = st.text_input("Temporary password", type="password")
                new_role = st.selectbox("Role", ["member", "reviewer", "admin"], index=0)
                if st.form_submit_button("Create user"):
                    try:
                        create_user(tenant_id=tenant_id, email=new_email, password=new_pwd, role=new_role)
                        st.success("User created.")
                    except Exception as e:
                        show_busy_error(e, context="Create user")

            users = list_users(tenant_id=tenant_id)
            if users:
                st.markdown("#### Users")
                selected_uid = st.selectbox(
                    "Select user",
                    options=[int(u["id"]) for u in users],
                    format_func=lambda uid: next(
                        (
                            f'{u["email"]} ({u["role"]}, {"active" if u["is_active"] else "disabled"})'
                            for u in users
                            if int(u["id"]) == int(uid)
                        ),
                        str(uid),
                    ),
                )
                selected_user = next((u for u in users if int(u["id"]) == int(selected_uid)), None)
                if selected_user:
                    col_u1, col_u2 = st.columns(2)
                    with col_u1:
                        toggle_label = "Disable user" if selected_user["is_active"] else "Enable user"
                        if st.button(toggle_label):
                            try:
                                set_user_active(
                                    tenant_id=tenant_id,
                                    user_id=int(selected_uid),
                                    is_active=not selected_user["is_active"],
                                )
                                st.success("Updated.")
                                st.rerun()
                            except Exception as e:
                                show_busy_error(e, context="Toggle user active")
                    with col_u2:
                        reset_pwd = st.text_input("Reset password", type="password", key=f"reset_pwd_{selected_uid}")
                        if st.button("Set new password"):
                            try:
                                reset_user_password(
                                    tenant_id=tenant_id,
                                    user_id=int(selected_uid),
                                    new_password=reset_pwd,
                                )
                                st.success("Password updated.")
                            except Exception as e:
                                show_busy_error(e, context="Reset password")

        with st.expander("Admin (Quotas)", expanded=False):
            st.caption("Set per-tenant daily limits. Set to 0 to disable a feature.")
            q = get_or_create_quota(tenant_id=tenant_id)
            with st.form("set_quotas"):
                gen_q = st.number_input(
                    "Generations per day", min_value=0, max_value=100000, value=int(q.get("generations_per_day", 0))
                )
                tr_q = st.number_input(
                    "Transcriptions per day", min_value=0, max_value=100000, value=int(q.get("transcriptions_per_day", 0))
                )
                vi_q = st.number_input(
                    "Vision analyses per day", min_value=0, max_value=100000, value=int(q.get("vision_analyses_per_day", 0))
                )
                if st.form_submit_button("Save quotas"):
                    try:
                        set_quota(
                            tenant_id=tenant_id,
                            generations_per_day=int(gen_q),
                            transcriptions_per_day=int(tr_q),
                            vision_analyses_per_day=int(vi_q),
                        )
                        st.success("Quotas updated.")
                        st.rerun()
                    except Exception as e:
                        show_busy_error(e, context="Set quotas")

st.session_state.company_manual_docs = load_manual_docs_with_chunks(
    tenant_id=tenant_id,
    manual_doc_ids=list(st.session_state.get("manual_multiselect_ids") or []),
)

template_name = str(st.session_state.get("gen_template") or "IT SOP")
strictness = str(st.session_state.get("gen_strictness") or "Detailed")
audience = str(st.session_state.get("gen_audience") or "")
tools_used = str(st.session_state.get("gen_tools_used") or "")
_gen_comp = st.session_state.get("gen_compliance")
compliance_standard = "" if _gen_comp in (None, "", "None") else str(_gen_comp)
tone = str(st.session_state.get("gen_tone") or "Professional")
include_definitions = bool(st.session_state.get("gen_inc_definitions", True))
include_safety_compliance = bool(st.session_state.get("gen_inc_safety", True))
include_records = bool(st.session_state.get("gen_inc_records", True))
include_checklist = bool(st.session_state.get("gen_inc_checklist", True))
rag_top_k = int(st.session_state.get("rag_top_k") or 6)

temperature = 0.35
model = "llama-3.1-8b-instant"

if "notes" not in st.session_state:
    st.session_state.notes = ""

api_key = get_groq_api_key()
if not api_key:
    st.warning(
        "Set `GROQ_API_KEY` in `.streamlit/secrets.toml` (see `.streamlit/secrets.toml.example`) "
        "or as the environment variable `GROQ_API_KEY` to generate SOPs."
    )

if active_page == "Generator":
    st.title(str(_brand.get("app_name") or DEFAULT_BRANDING["app_name"]))
    st.markdown("### What do you want to document today?")
    _tag = header_tagline(_brand)
    if _tag:
        st.caption(_tag)

    generate_btn = False
    tab_text, tab_voice, tab_vision = st.tabs(
        ["📝 Text / Notes", "🎤 Voice Recording", "🖼️ Vision / Image"]
    )

    with tab_text:
        with st.container(border=True):
            st.markdown("#### 📄 Paste your raw notes")
            st.text_area(
                "Input notes / raw text",
                key="notes",
                height=250,
                placeholder="Paste rough notes, emails, or chat logs here…",
                label_visibility="collapsed",
            )
            hc1, hc2, hc3 = st.columns([1, 2, 1])
            with hc2:
                generate_btn = st.button(
                    "✨ Generate Standard Operating Procedure",
                    type="primary",
                    use_container_width=True,
                    disabled=not api_key,
                    key="hero_generate_sop",
                )

    with tab_voice:
        with st.container(border=True):
            st.markdown("#### 🎙️ Talk through the process")
            st.info("Describe the task out loud. We will transcribe it and fill your **Text / Notes** tab.")
            audio_rec = None
            if hasattr(st, "audio_input"):
                audio_rec = st.audio_input("Record audio", key="dash_audio_input")
            st.caption("Or upload an audio file")
            audio_file = st.file_uploader(
                "Upload audio",
                type=["wav", "mp3", "m4a", "aac", "flac", "ogg", "webm"],
                accept_multiple_files=False,
                key="dash_audio_upload",
            )
            stt_model = st.selectbox(
                "Speech-to-text model",
                ["whisper-large-v3-turbo", "whisper-large-v3"],
                index=0,
                key="dash_stt_model",
            )
            stt_language = st.text_input(
                "Language (optional, ISO-639-1)", value="", placeholder="e.g., en", key="dash_stt_lang"
            )
            _has_audio = (audio_rec is not None) or (audio_file is not None)
            if st.button("Transcribe audio", disabled=(not api_key or not _has_audio), key="dash_transcribe"):
                try:
                    ok, msg = check_and_consume_quota(tenant_id=TENANT_ID, action="transcribe", amount=1)
                    if not ok:
                        st.error(msg)
                        st.stop()
                    if audio_rec is not None:
                        audio_bytes = audio_rec.getvalue()
                        audio_name = getattr(audio_rec, "name", None) or "recording.webm"
                    else:
                        audio_bytes = audio_file.getvalue()
                        audio_name = audio_file.name
                    file_sha = hashlib.sha256(audio_bytes).hexdigest()
                    with st.spinner("Transcribing..."):
                        transcript = transcribe_audio_cached(
                            api_key=api_key,
                            model=stt_model,
                            file_name=audio_name,
                            file_sha256=file_sha,
                            audio_bytes=audio_bytes,
                            language=stt_language.strip(),
                        )
                    if transcript:
                        st.session_state.notes = transcript
                        st.success("Transcription complete. Switch to **Text / Notes** to review or edit.")
                    else:
                        st.error("Transcription returned empty text.")
                except Exception as e:
                    show_busy_error(e, context="Transcribe audio")

    with tab_vision:
        with st.container(border=True):
            st.markdown("#### 📸 Analyze a screenshot")
            st.caption("Upload an interface screenshot or photo; we extract structured notes into **Text / Notes**.")
            image_file = st.file_uploader(
                "Upload interface screenshot",
                type=["png", "jpg", "jpeg", "webp"],
                accept_multiple_files=False,
                key="dash_vision_upload",
            )
            vision_model = st.selectbox(
                "Vision model",
                ["meta-llama/llama-4-scout-17b-16e-instruct"],
                index=0,
                key="dash_vision_model",
            )
            if image_file is not None:
                st.image(image_file, caption=image_file.name, use_container_width=True)
            if st.button("Analyze image", disabled=(not api_key or image_file is None), key="dash_analyze_image"):
                try:
                    ok, msg = check_and_consume_quota(tenant_id=TENANT_ID, action="vision", amount=1)
                    if not ok:
                        st.error(msg)
                        st.stop()
                    image_bytes = image_file.getvalue()
                    file_sha = hashlib.sha256(image_bytes).hexdigest()
                    mime_type = image_file.type or "image/png"
                    image_b64 = base64.b64encode(image_bytes).decode("utf-8")
                    with st.spinner("Analyzing image..."):
                        extracted_notes = analyze_image_to_notes_cached(
                            api_key=api_key,
                            model=vision_model,
                            file_sha256=file_sha,
                            mime_type=mime_type,
                            image_b64=image_b64,
                        )
                    if extracted_notes:
                        st.session_state.notes = extracted_notes
                        st.success("Image analysis complete. Switch to **Text / Notes** to review or edit.")
                    else:
                        st.error("Image analysis returned empty text.")
                except Exception as e:
                    show_busy_error(e, context="Analyze image")

    with st.expander("Recent SOPs (last 5)", expanded=False):
        history_items = load_history(tenant_id=tenant_id, user_id=user_id, role=user_role, limit=5)
        if history_items:
            options = []
            for i, it in enumerate(history_items):
                ts = (it.get("ts") or "")[:19].replace("T", " ")
                label = it.get("label") or it.get("template_name") or "SOP"
                status = (it.get("status") or "draft").replace("_", " ")
                options.append(f"{i+1}. {label} [{status}] — {ts}")
            selected = st.selectbox("Saved SOPs", options, index=0, key="dash_hist_pick")
            sel_idx = int(selected.split(".")[0]) - 1
            selected_item = history_items[sel_idx]
            if st.button("Load into editor", key="dash_hist_load"):
                st.session_state.current_sop_text = selected_item.get("sop_text", "") or ""
                st.session_state.last_inferred_topic = selected_item.get("label", "SOP") or "SOP"
                st.session_state.current_sop_doc_id = selected_item.get("sop_doc_id")
                st.success("Loaded into editor.")
        else:
            st.caption("No saved SOPs yet.")

    notes = str(st.session_state.get("notes") or "")

    if generate_btn:
        sop_text = ""
        if not notes.strip():
            st.error("Please paste your notes (or a transcript) first.")
        else:
            with st.spinner("Writing SOP..."):
                try:
                    ok, msg = check_and_consume_quota(tenant_id=TENANT_ID, action="generate", amount=1)
                    if not ok:
                        st.error(msg)
                        st.stop()
                    inferred_topic = f"{template_name} SOP"
                    # Build company brain context for this generation (stored in session_state).
                    company_docs = st.session_state.get("company_manual_docs", []) or []
                    st.session_state.company_rules_context = ""
                    if company_docs:
                        query = f"{template_name}\n{inferred_topic}\n{audience}\n{tools_used}\n{compliance_standard}\n{notes}"
                        snippets = retrieve_company_snippets(query=query, docs=company_docs, top_k=int(rag_top_k))
                        if snippets:
                            ctx_lines = []
                            for s in snippets:
                                ctx_lines.append(
                                    f"- ({s['doc']} #chunk{s['chunk_index']}) {s['text']}"
                                )
                            st.session_state.company_rules_context = "\n".join(ctx_lines)

                    prompt = build_prompt_for_template(
                        template_name,
                        inferred_topic,
                        notes,
                        audience=audience.strip() or "General staff",
                        tools_used=tools_used.strip(),
                        compliance_standard=compliance_standard.strip(),
                        strictness=strictness,
                        tone=tone,
                        include_definitions=include_definitions,
                        include_safety_compliance=include_safety_compliance,
                        include_records=include_records,
                        include_checklist=include_checklist,
                    )
                    sop_text = generate_sop_cached(
                        api_key=api_key,
                        model=model,
                        temperature=float(temperature),
                        prompt=prompt,
                    )
                    st.session_state.last_sop_text = sop_text
                    st.session_state.last_inferred_topic = inferred_topic
                    # Always set the "current" SOP so it persists across any subsequent button clicks.
                    st.session_state.current_sop_text = sop_text

                    add_to_history(
                        tenant_id=TENANT_ID,
                        user_id=USER_ID,
                        role=USER_ROLE,
                        entry={
                            "ts": datetime.now(timezone.utc).isoformat(),
                            "label": inferred_topic,
                            "template_name": template_name,
                            "source": "generated",
                            "sop_text": sop_text,
                            "sop_sha256": sop_fingerprint(sop_text),
                            "sop_doc_id": st.session_state.get("current_sop_doc_id"),
                        },
                    )
                except Exception as e:
                    show_busy_error(e, context="Generate SOP")
                    sop_text = ""

            if sop_text:
                # Enforce strict Markdown structure. If the model drifts, auto-fix it once.
                if not is_valid_sop_markdown(sop_text):
                    try:
                        sop_text = format_sop_to_required_markdown_cached(
                            api_key=api_key,
                            model=model,
                            sop_text=sop_text,
                        )
                    except Exception as e:
                        show_busy_error(e, context="Format SOP to required Markdown")

                st.session_state.last_sop_text = sop_text
                st.session_state.current_sop_text = sop_text
                st.success("SOP generated. See 'Current SOP' below.")

elif active_page == "Library":
    st.title("📂 SOP Library")
    if str(USER_ROLE or "").strip().lower() == "member":
        st.caption("Showing **your** SOPs only (saved under your account). Admins and reviewers see all workspace SOPs.")
    q = st.text_input("Search", value="", placeholder="Search by title or template…")
    docs = list_sop_docs(tenant_id=TENANT_ID, user_id=USER_ID, role=USER_ROLE, query=q, limit=200)
    if not docs:
        st.info("No SOPs yet.")
    else:
        labels = [
            f'{d["title"]} · {d["template_name"]} · v{d["latest_version"]} · {(d["latest_status"] or "draft").replace("_"," ")}'
            for d in docs
        ]
        sel = st.selectbox("Open SOP", options=list(range(len(docs))), format_func=lambda i: labels[int(i)], index=0)
        doc = docs[int(sel)]
        st.session_state.current_sop_doc_id = doc["id"]
        st.session_state.last_inferred_topic = doc["title"]
        st.session_state.current_sop_text = doc["latest_text"] or ""
        st.markdown(f'**Latest:** v{doc["latest_version"]} · {(doc["latest_status"] or "draft").replace("_"," ")}')
        st.markdown(st.session_state.current_sop_text or "")

        with st.expander("Versions", expanded=False):
            versions = load_doc_versions(
                tenant_id=TENANT_ID, sop_doc_id=int(doc["id"]), user_id=USER_ID, role=USER_ROLE, limit=50
            )
            v_labels = [
                f'v{v["version"]} · {str(v["status"]).replace("_"," ")} · {(v["ts"] or "")[:19].replace("T"," ")} · {v.get("source","")}'
                for v in versions
            ]
            v_a = st.selectbox("Version A", options=list(range(len(versions))), format_func=lambda i: v_labels[int(i)], index=0, key="lib_diff_a")
            v_b = st.selectbox("Version B", options=list(range(len(versions))), format_func=lambda i: v_labels[int(i)], index=min(1, len(versions) - 1), key="lib_diff_b")
            if st.button("Show diff", key="lib_show_diff"):
                st.code(diff_text(versions[int(v_a)]["sop_text"], versions[int(v_b)]["sop_text"]))


# --- Persistent SOP display (state management) ---
current_sop = (st.session_state.get("current_sop_text") or "").strip()
if current_sop:
    st.divider()
    st.subheader("Current SOP")
    st.markdown(current_sop)

    inferred_topic = str(st.session_state.get("last_inferred_topic") or "sop")
    safe_name = "".join(c for c in inferred_topic.strip() if c.isalnum() or c in (" ", "-", "_")).strip() or "sop"
    current_pdf = b""
    current_docx = b""
    try:
        current_pdf = create_pdf_bytes(current_sop)
    except Exception as e:
        log_exception(e, context="Prepare current SOP PDF")
        show_busy_error(e, context="Prepare current SOP PDF")
    try:
        current_docx = create_docx_bytes(safe_name, current_sop)
    except Exception as e:
        log_exception(e, context="Prepare current SOP DOCX")
        show_busy_error(e, context="Prepare current SOP DOCX")

    with st.expander("Edit SOP (client-ready)", expanded=False):
        edited_current = st.text_area(
            "Edit SOP text",
            value=current_sop,
            height=320,
            key=f"editor_current_{sop_fingerprint(current_sop)}",
        )
        edit_note = st.text_input(
            "Change note (optional)",
            value="",
            placeholder="e.g., Clarified responsibilities, added exception handling",
            key=f"edit_note_current_{sop_fingerprint(current_sop)}",
        )
        col_ec1, col_ec2 = st.columns(2)
        with col_ec1:
            if st.button("Save edits", key=f"save_current_{sop_fingerprint(current_sop)}"):
                try:
                    st.session_state.current_sop_text = edited_current
                    save_current_edits_as_new_version(
                        tenant_id=TENANT_ID,
                        user_id=USER_ID,
                        role=USER_ROLE,
                        template_name=template_name,
                        label=str(st.session_state.get("last_inferred_topic") or "SOP"),
                        sop_text=edited_current,
                        source="edited",
                        change_note=edit_note.strip(),
                    )
                    st.success("Edits saved as a new version.")
                except Exception as e:
                    show_busy_error(e, context="Save current edits as version")
        with col_ec2:
            if st.button("Reset to last generated", key="reset_current_to_last"):
                st.session_state.current_sop_text = st.session_state.get("last_sop_text", "") or current_sop
                st.success("Reset.")

    col_dl1, col_dl2 = st.columns(2)
    with col_dl1:
        if current_pdf:
            st.download_button(
                "Download PDF",
                data=current_pdf,
                file_name=f"{safe_name}.pdf",
                mime="application/pdf",
                key=f"dl_pdf_current_{sop_fingerprint(current_sop)}",
            )
    with col_dl2:
        if current_docx:
            st.download_button(
                "Download DOCX",
                data=current_docx,
                file_name=f"{safe_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_docx_current_{sop_fingerprint(current_sop)}",
            )


# --- Step 2: Quality pass (Review & Fix) ---
last_sop = st.session_state.get("last_sop_text", "")
if api_key and last_sop:
    st.divider()
    st.subheader("Review & Fix SOP (Quality pass)")
    st.caption("Runs an editor pass to fix gaps, unclear steps, and missing roles/records.")

    do_review = st.button("Review & Fix SOP", type="secondary")
    if do_review:
        with st.spinner("Reviewing and improving SOP..."):
            try:
                fixed = review_and_fix_sop_cached(
                    api_key=api_key,
                    model=model,
                    temperature=min(float(temperature), 0.4),
                    sop_text=last_sop,
                    strictness=strictness,
                    tone=tone,
                    compliance_standard=compliance_standard.strip(),
                )
                st.session_state.last_fixed_sop_text = fixed

                inferred_topic = st.session_state.get("last_inferred_topic", "SOP")
                add_to_history(
                    tenant_id=TENANT_ID,
                    user_id=USER_ID,
                    role=USER_ROLE,
                    entry={
                        "ts": datetime.now(timezone.utc).isoformat(),
                        "label": f"{inferred_topic} (revised)",
                        "template_name": template_name,
                        "source": "revised",
                        "sop_text": fixed,
                        "sop_sha256": sop_fingerprint(fixed),
                    },
                )
            except Exception as e:
                show_busy_error()

    fixed_sop = st.session_state.get("last_fixed_sop_text", "")
    if fixed_sop:
        st.subheader("Revised SOP")
        st.markdown(fixed_sop)

        # Make revised SOP the current editable SOP by default.
        st.session_state.current_sop_text = fixed_sop

        inferred_topic = st.session_state.get("last_inferred_topic", "SOP")
        safe_name = (
            "".join(c for c in str(inferred_topic).strip() if c.isalnum() or c in (" ", "-", "_")).strip()
            or "sop"
        )

        with st.expander("Interactive Step Editor (edit revised SOP)", expanded=False):
            st.caption("Edit the revised SOP here, then download the edited version.")
            edited_rev = st.text_area(
                "Edit revised SOP text",
                value=st.session_state.current_sop_text,
                height=320,
                key=f"editor_rev_{sop_fingerprint(fixed_sop)}",
            )
            edit_note_rev = st.text_input(
                "Change note (optional)",
                value="",
                placeholder="e.g., Tightened steps, added records/evidence",
                key=f"edit_note_rev_{sop_fingerprint(fixed_sop)}",
            )
            col_r1, col_r2 = st.columns(2)
            with col_r1:
                if st.button("Save revised edits", key=f"save_rev_{sop_fingerprint(fixed_sop)}"):
                    try:
                        st.session_state.current_sop_text = edited_rev
                        save_current_edits_as_new_version(
                            tenant_id=TENANT_ID,
                            user_id=USER_ID,
                            role=USER_ROLE,
                            template_name=template_name,
                            label=str(st.session_state.get("last_inferred_topic") or "SOP"),
                            sop_text=edited_rev,
                            source="edited",
                            change_note=edit_note_rev.strip(),
                        )
                        st.success("Edits saved as a new version. Downloads will use the edited SOP.")
                    except Exception as e:
                        show_busy_error(e, context="Save revised edits as version")
            with col_r2:
                if st.button("Reset to revised", key=f"reset_rev_{sop_fingerprint(fixed_sop)}"):
                    st.session_state.current_sop_text = fixed_sop
                    st.success("Reset to the revised SOP.")

        sop_for_download = st.session_state.get("current_sop_text") or fixed_sop
        pdf_bytes = b""
        docx_bytes = b""
        try:
            pdf_bytes = create_pdf_bytes(sop_for_download)
        except Exception as e:
            log_exception(e, context="Prepare revised SOP PDF")
            show_busy_error(e, context="Prepare revised SOP PDF")
        try:
            docx_bytes = create_docx_bytes(safe_name, sop_for_download)
        except Exception as e:
            log_exception(e, context="Prepare revised SOP DOCX")
            show_busy_error(e, context="Prepare revised SOP DOCX")

        col_c, col_d = st.columns(2)
        with col_c:
            if pdf_bytes:
                st.download_button(
                    "Download Revised PDF",
                    data=pdf_bytes,
                    file_name=f"{safe_name}-revised.pdf",
                    mime="application/pdf",
                    key=f"dl_pdf_rev_{sop_fingerprint(sop_for_download)}",
                )
        with col_d:
            if docx_bytes:
                st.download_button(
                    "Download Revised DOCX",
                    data=docx_bytes,
                    file_name=f"{safe_name}-revised.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_docx_rev_{sop_fingerprint(sop_for_download)}",
                )

        st.markdown("### Rate the revised SOP")
        rating2 = st.radio(
            "Was the revised SOP helpful?",
            ["👍 Thumbs Up", "👎 Thumbs Down"],
            horizontal=True,
            key=f"rating_rev_{sop_fingerprint(fixed_sop)}",
        )
        reason2 = ""
        if rating2.startswith("👎"):
            reason2 = st.text_area(
                "What should still be improved?",
                placeholder="Be specific about what’s missing or unclear.",
                key=f"reason_rev_{sop_fingerprint(fixed_sop)}",
            )
        if st.button("Submit revised feedback", key=f"submit_rev_{sop_fingerprint(fixed_sop)}"):
            append_feedback(
                tenant_id=TENANT_ID,
                entry={
                    "ts": datetime.now(timezone.utc).isoformat(),
                    "rating": "up" if rating2.startswith("👍") else "down",
                    "reason": reason2.strip(),
                    "template_name": template_name,
                    "strictness": strictness,
                    "tone": tone,
                    "compliance_standard": compliance_standard or "",
                    "audience": (audience or "").strip(),
                    "tools_used": (tools_used or "").strip(),
                    "include_definitions": bool(include_definitions),
                    "include_safety_compliance": bool(include_safety_compliance),
                    "include_records": bool(include_records),
                    "include_checklist": bool(include_checklist),
                    "model": model,
                    "temperature": float(temperature),
                    "notes_chars": len((st.session_state.get("notes") or "").strip()),
                    "sop_sha256": sop_fingerprint(fixed_sop),
                    "source": "revised",
                }
            )
            st.success("Thanks — feedback saved.")


# --- Compliance Auditor (AI Critic) ---
candidate_sop_for_audit = st.session_state.get("last_fixed_sop_text") or st.session_state.get("last_sop_text") or ""
if api_key and candidate_sop_for_audit:
    st.divider()
    st.subheader("Compliance Auditor (AI Critic)")
    st.caption("A second AI pass that critiques the SOP for gaps, risks, and missing compliance items.")

    run_audit = st.button("Review SOP (Auditor)")
    if run_audit:
        with st.spinner("Auditing SOP..."):
            try:
                audit = compliance_audit_cached(
                    api_key=api_key,
                    model=model,
                    temperature=0.2,
                    sop_text=candidate_sop_for_audit,
                    template_name=template_name,
                    strictness=strictness,
                    tone=tone,
                    compliance_standard=compliance_standard.strip(),
                    company_rules_context=st.session_state.get("company_rules_context", "").strip(),
                )
                st.session_state.last_audit_text = audit
            except Exception as e:
                show_busy_error()

    audit_text = st.session_state.get("last_audit_text", "")
    if audit_text:
        st.markdown(audit_text)


# --- Visual Flowchart ---
candidate_sop_for_flowchart = st.session_state.get("last_fixed_sop_text") or st.session_state.get("last_sop_text") or ""
if api_key and candidate_sop_for_flowchart:
    st.divider()
    st.subheader("Visual flowchart")
    st.caption("Generates a flowchart from the latest SOP (revised if available).")

    gen_chart = st.button("Generate Flowchart")
    if gen_chart:
        with st.spinner("Generating flowchart..."):
            try:
                mermaid_code = generate_flowchart_mermaid_cached(
                    api_key=api_key,
                    model=model,
                    temperature=0.2,
                    sop_text=candidate_sop_for_flowchart,
                )
                st.session_state.last_mermaid_flowchart = mermaid_code
            except Exception as e:
                show_busy_error()

    mermaid_code = st.session_state.get("last_mermaid_flowchart", "")
    if mermaid_code:
        render_mermaid(mermaid_code, height_px=700)
        st.download_button(
            "Download Flowchart (Mermaid)",
            data=mermaid_code.encode("utf-8"),
            file_name="sop-flowchart.mmd",
            mime="text/plain",
        )
